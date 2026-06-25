import streamlit as st
import openpyxl
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import zipfile
import io
import base64
import os
import requests
from datetime import datetime

def read_gsheet_as_rows(url):
    """Đọc Google Sheet public từ URL, trả về list of rows (list of values)."""
    import re
    match = re.search(r"/spreadsheets/d/([a-zA-Z0-9-_]+)", url)
    if not match:
        raise ValueError("URL không hợp lệ, không tìm thấy Sheet ID.")
    sheet_id = match.group(1)
    
    # Lấy gid nếu có
    gid_match = re.search(r"gid=(\d+)", url)
    gid = gid_match.group(1) if gid_match else "0"
    
    export_url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=xlsx&gid={gid}"
    response = requests.get(export_url, timeout=15)
    if response.status_code != 200:
        raise ValueError(f"Không tải được file. Status: {response.status_code}. Sheet có thể chưa public.")
    return response.content  # trả về bytes giống uploaded file
# ── Nút góp ý cố định góc dưới phải ──
st.markdown("""
    <style>
    .gop-y-btn {
        position: fixed;
        bottom: 30px;
        right: 30px;
        z-index: 9999;
    }
    .gop-y-btn a {
        background-color: #2B547E;
        color: white !important;
        padding: 12px 20px;
        border-radius: 25px;
        text-decoration: none;
        font-weight: bold;
        font-size: 15px;
        box-shadow: 0 4px 12px rgba(0,0,0,0.3);
    }
    .gop-y-btn a:hover {
        background-color: #1F4E79;
    }
    </style>
    <div class="gop-y-btn">
        <a href="https://forms.gle/ofYqU5vrehs8s5kr8" target="_blank">📝 Góp ý</a>
    </div>
""", unsafe_allow_html=True)

# ─── Cấu hình trang ───────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Tạo Phiếu Tự động - Phòng Mua Hàng",
    page_icon="📦",
    layout="wide"
)

st.markdown("""
<style>
/* ── Ẩn header mặc định của Streamlit ── */
#MainMenu, header, footer { visibility: hidden; }
.block-container { padding-top: 0 !important; }

/* ── Navbar ── */
.navbar {
    display: flex;
    align-items: center;
    justify-content: space-between;
    background: #ffffff;
    border-bottom: 2px solid #e8edf2;
    padding: 0 32px;
    height: 64px;
    position: sticky;
    top: 0;
    z-index: 999;
    box-shadow: 0 2px 12px rgba(43,84,126,0.08);
}
.navbar-left {
    display: flex;
    align-items: center;
    gap: 12px;
}
.navbar-logo {
    height: 38px;
    object-fit: contain;
}
.navbar-title {
    font-family: 'Segoe UI', Tahoma, sans-serif;
    font-size: 17px;
    font-weight: 700;
    color: #2B547E;
    letter-spacing: 0.3px;
    white-space: nowrap;
}
.navbar-subtitle {
    font-size: 11px;
    color: #8a9bb0;
    font-family: 'Segoe UI', Tahoma, sans-serif;
    margin-top: -2px;
}
.navbar-divider {
    width: 1px;
    height: 32px;
    background: #dde3ea;
    margin: 0 4px;
}

/* ── Nav tabs ── */
.nav-tabs {
    display: flex;
    gap: 4px;
    align-items: center;
    height: 100%;
}
.nav-tab {
    display: flex;
    align-items: center;
    gap: 6px;
    padding: 8px 18px;
    border-radius: 8px;
    font-size: 13.5px;
    font-weight: 600;
    font-family: 'Segoe UI', Tahoma, sans-serif;
    cursor: pointer;
    text-decoration: none;
    border: none;
    background: transparent;
    color: #5a7a9a;
    transition: all 0.18s ease;
    white-space: nowrap;
}
.nav-tab:hover {
    background: #f0f4f8;
    color: #2B547E;
}
.nav-tab.active {
    background: #2B547E;
    color: #ffffff;
    box-shadow: 0 2px 8px rgba(43,84,126,0.25);
}

/* ── Page content ── */
.page-content {
    max-width: 1100px;
    margin: 32px auto 0 auto;
    padding: 0 24px;
}
.page-header {
    margin-bottom: 24px;
}
.page-header h2 {
    font-family: 'Segoe UI', Tahoma, sans-serif;
    font-size: 26px;
    font-weight: 700;
    color: #1a3a5c;
    margin: 0 0 4px 0;
}
.page-header p {
    color: #7a93aa;
    font-size: 13px;
    margin: 0;
}
</style>
""", unsafe_allow_html=True)

# ── Navbar HTML ──
LOGO_URL_NAV = "https://static.ybox.vn/2023/7/4/1689816973012-LOGO.jpg"

if "active_page" not in st.session_state:
    st.session_state.active_page = "xuat_kho"

# Đọc query_params TRƯỚC khi render sidebar
_qp_early = st.query_params.get("page", "")
if _qp_early in ["xuat_kho", "chuyen_hang", "hang_phong_choi", "hang_game"]:
    st.session_state.active_page = _qp_early

# Render navbar
st.markdown(f"""
<div class="navbar">
    <div class="navbar-left">
        <img class="navbar-logo" src="{LOGO_URL_NAV}" />
        <div class="navbar-divider"></div>
        <div>
            <div class="navbar-title">MINH ĐỨC</div>
            <div class="navbar-subtitle">Hệ thống quản lý phiếu kho</div>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)

# ── Navigation ──
st.markdown("""
<style>
[data-testid="stSidebar"] { display: none !important; }
[data-testid="collapsedControl"] { display: none !important; }
.main .block-container { padding-left: 244px !important; padding-right: 32px !important; }

/* ── Sliding sidebar ── */
.slide-sidebar {
    position: fixed;
    top: 0; left: 0;
    height: 100vh;
    width: 56px;
    background: #1a3a5c;
    overflow: hidden;
    transition: width 0.3s cubic-bezier(.4,0,.2,1);
    z-index: 1000;
    box-shadow: 3px 0 16px rgba(0,0,0,0.18);
    display: flex;
    flex-direction: column;
    align-items: flex-start;
}
.slide-sidebar:hover {
    width: 220px;
}

/* Logo */
.slide-logo {
    display: flex;
    align-items: center;
    padding: 14px 10px 10px 10px;
    border-bottom: 1px solid rgba(255,255,255,0.1);
    min-width: 220px;
    height: 64px;
}
.slide-logo img {
    width: 36px; height: 36px;
    border-radius: 6px;
    object-fit: contain;
    background: white;
    flex-shrink: 0;
}
.slide-logo span {
    font-family: 'Segoe UI', Tahoma, sans-serif;
    font-size: 13px;
    font-weight: 700;
    color: white;
    margin-left: 12px;
    white-space: nowrap;
    opacity: 0;
    transition: opacity 0.2s ease 0.1s;
}
.slide-sidebar:hover .slide-logo span { opacity: 1; }

/* Menu items */
.slide-menu {
    margin-top: 16px;
    width: 100%;
    flex: 1;
}
.slide-item {
    display: flex;
    align-items: center;
    padding: 12px 16px;
    cursor: pointer;
    border-radius: 0 10px 10px 0;
    margin: 4px 8px 4px 0;
    min-width: 220px;
    transition: background 0.2s;
    text-decoration: none;
}
.slide-item:hover { background: rgba(255,255,255,0.12); }
.slide-item.active { background: #2B547E; }
.slide-icon {
    font-size: 20px;
    flex-shrink: 0;
    width: 28px;
    text-align: center;
}
.slide-label {
    font-family: 'Segoe UI', Tahoma, sans-serif;
    font-size: 13.5px;
    font-weight: 600;
    color: white;
    margin-left: 14px;
    white-space: nowrap;
    opacity: 0;
    transition: opacity 0.2s ease 0.1s;
}
.slide-sidebar:hover .slide-label { opacity: 1; }

/* Footer */
.slide-footer {
    padding: 16px 10px;
    border-top: 1px solid rgba(255,255,255,0.1);
    min-width: 220px;
}
.slide-footer span {
    font-size: 10px;
    color: rgba(255,255,255,0.4);
    white-space: nowrap;
    opacity: 0;
    transition: opacity 0.2s ease 0.1s;
}
.slide-sidebar:hover .slide-footer span { opacity: 1; }
</style>
""", unsafe_allow_html=True)

# Xác định active
_p = st.session_state.get("active_page", "xuat_kho")
_cls_xk = "slide-item active" if _p == "xuat_kho"    else "slide-item"
_cls_ch = "slide-item active" if _p == "chuyen_hang" else "slide-item"
_cls_hpc = "slide-item active" if _p == "hang_phong_choi" else "slide-item"
_cls_game = "slide-item active" if _p == "hang_game" else "slide-item"

st.markdown(f"""
<div class="slide-sidebar">
    <div class="slide-logo">
        <img src="https://i.ibb.co/wZpd32qL/149a4595-c9c7-487f-ab75-be439ed2fa85.png"/>
        <span>PHÒNG MUA HÀNG</span>
    </div>
    <div class="slide-menu">
        <a class="{_cls_xk}" href="?page=xuat_kho" target="_self">
            <span class="slide-icon">🏭</span>
            <span class="slide-label">Phiếu Xuất Kho</span>
        </a>
        <a class="{_cls_ch}" href="?page=chuyen_hang" target="_self">
            <span class="slide-icon">🚚</span>
            <span class="slide-label">Phiếu Chuyển Hàng</span>
        </a>
        <a class="{_cls_hpc}" href="?page=hang_phong_choi" target="_self">
            <span class="slide-icon">🎮</span>
            <span class="slide-label">Hàng Phòng Chơi</span>
        </a>
        <a class="{_cls_game}" href="?page=hang_game" target="_self">
            <span class="slide-icon">🕹️</span>
            <span class="slide-label">Hàng Game</span>
        </a>
    </div>
    <div class="slide-footer">
        <span>Phòng Mua Hàng · Minh Đức</span>
    </div>
</div>
""", unsafe_allow_html=True)

page = st.session_state.get("active_page", "xuat_kho")


# ─── Logo URL ─────────────────────────────────────────────────────────────────
LOGO_URL = "https://static.ybox.vn/2023/7/4/1689816973012-LOGO.jpg"

def load_logo_from_url(url):
    try:
        response = requests.get(url, timeout=5)
        response.raise_for_status()
        return response.content
    except Exception:
        return None

# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_border(cell, color="CCCCCC"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:color"), color)
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def fmt_number(n):
    try:
        return f"{int(n):,}"
    except:
        return "0"

def make_ma_phieu(dt, index=0):
    base = dt.strftime('%d%m%y%H%M%S')
    return f"PXK-{base}-{index+1:03d}"
#__________________________
def get_gsheet():
    import gspread
    from google.oauth2.service_account import Credentials
    scope = [
        "https://spreadsheets.google.com/feeds",
        "https://www.googleapis.com/auth/drive"
    ]
    creds = Credentials.from_service_account_info(
        st.secrets["gcp_service_account"], scopes=scope
    )
    client = gspread.authorize(creds)
    sheet = client.open_by_key(st.secrets["SHEET_ID"]).sheet1
    return sheet

def merge_items_by_ma_hang(items):
    from collections import OrderedDict
    merged = OrderedDict()
    for item in items:
        ma  = item.get("ma_hang", "").strip()
        ten = item.get("ten_hang", "").strip()
        sl  = item.get("sl_chuyen", item.get("Số đặt", 0)) or 0
        key = ma if ma else f"__nocode__{ten}"
        if key not in merged:
            merged[key] = {
                "ma_hang":   ma,
                "ten_hang":  ten,
                "dvt":       item.get("dvt", ""),
                "sl_chuyen": sl,
                "_ten_list": [ten] if ten else [],
                "_chi_tiet": [f"{ten}: {sl}"] if ten else [],
            }
        else:
            merged[key]["sl_chuyen"] += sl
            if ten and ten not in merged[key]["_ten_list"]:
                merged[key]["_ten_list"].append(ten)
                merged[key]["_chi_tiet"].append(f"{ten}: {sl}")
            merged[key]["ten_hang"] = " / ".join(merged[key]["_ten_list"])
    result = []
    for v in merged.values():
        v.pop("_ten_list", None)
        # Chỉ ghi chú nếu có gộp (>1 sản phẩm)
        chi_tiet = v.pop("_chi_tiet", [])
        v["ghi_chu_gop"] = " | ".join(chi_tiet) if len(chi_tiet) > 1 else ""
        result.append(v)
    return result

def log_phieu_chuyen_to_sheet(ma_phieu, ngay_chuyen, noi_chuyen, tt_nhan, items):
    try:
        sheet = get_gsheet()
        items = merge_items_by_ma_hang(items)   # ← THÊM DÒNG NÀY
        rows = []
        for item in items:
            rows.append([
                ma_phieu, ngay_chuyen, noi_chuyen, tt_nhan,
                item.get("ten_hang", ""), item.get("ma_hang", ""),
                item.get("sl_chuyen", 0), "", item.get("ghi_chu_gop", ""),
            ])

        sheet.insert_rows(rows, row=2, value_input_option="USER_ENTERED")

        # Gom tất cả format vào 1 batch request duy nhất
        spreadsheet = sheet.spreadsheet
        sheet_id = sheet.id
        requests = [
            # Reset toàn bộ về trắng
            {
                "repeatCell": {
                    "range": {"sheetId": sheet_id, "startRowIndex": 1, "endRowIndex": 1 + len(rows), "startColumnIndex": 0, "endColumnIndex": 9},
                    "cell": {"userEnteredFormat": {"backgroundColor": {"red": 1, "green": 1, "blue": 1}, "textFormat": {"bold": False}}},
                    "fields": "userEnteredFormat(backgroundColor,textFormat)"
                }
            },
            # Highlight dòng đầu màu xanh
            {
                "repeatCell": {
                    "range": {"sheetId": sheet_id, "startRowIndex": 1, "endRowIndex": 2, "startColumnIndex": 0, "endColumnIndex": 9},
                    "cell": {"userEnteredFormat": {"backgroundColor": {"red": 0.68, "green": 0.85, "blue": 0.90}, "textFormat": {"bold": True}}},
                    "fields": "userEnteredFormat(backgroundColor,textFormat)"
                }
            }
        ]
        spreadsheet.batch_update({"requests": requests})
        return True
    except Exception as e:
        st.warning(f"⚠️ Không thể ghi Google Sheet: {e}")
        return False
# ─── Đọc Excel ────────────────────────────────────────────────────────────────
def parse_excel(file_bytes):
    """
    Đọc file Excel định dạng chuẩn:
      Hàng 1: tên trung tâm (cột 7, 9, 11... xen kẽ None)
      Hàng 2: NCC | Mã hàng | Tên hàng | ĐVT | Đơn giá | Số đặt | Thành Tiền | Số đặt | ...
      Hàng 3: Tổng Cộng (bỏ qua)
      Hàng 4+: dữ liệu
    """
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    all_data = {}

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if len(rows) < 4:
            continue

        row1 = rows[0]  # tên trung tâm
        row2 = rows[1]  # sub-header: NCC, Mã hàng, Tên hàng, ĐVT, Đơn giá, Số đặt...

        # ── Tìm cột cố định từ hàng 2 ────────────────────────────────────
        col_ncc      = next((i for i, v in enumerate(row2) if v and str(v).strip() == "NCC"), 0)
        col_ma_hang  = next((i for i, v in enumerate(row2) if v and "mã hàng" in str(v).strip().lower()), 1)
        col_ten_hang = next((i for i, v in enumerate(row2) if v and "tên hàng" in str(v).strip().lower()), 2)
        col_dvt      = next((i for i, v in enumerate(row2) if v and str(v).strip().upper() == "ĐVT"), 3)
        col_gia_cost = next((i for i, v in enumerate(row2) if v and "đơn giá" in str(v).strip().lower()), 4)

        # ── Tìm trung tâm từ hàng 1, map sang cột "Số đặt" ở hàng 2 ─────
        # Mỗi trung tâm ở hàng 1 tại cột i, "Số đặt" của nó ở cùng cột i hàng 2
        branches = {}
        for i, val in enumerate(row1):
            if not val or not str(val).strip():
                continue
            if i < col_gia_cost:
                continue
            # Bỏ qua cột Tổng order
            if "tổng" in str(val).strip().lower():
                continue
            # Xác nhận cột này là "Số đặt" ở hàng 2
            if i < len(row2) and row2[i] and "đặt" in str(row2[i]).lower():
                branches[str(val).strip()] = i

        # ── Đọc dữ liệu từ hàng 4 (index 3), bỏ hàng Tổng Cộng ─────────
        products = []
        for row in rows[3:]:
            if not row or len(row) <= col_ten_hang:
                continue
            ten_hang = row[col_ten_hang]
            if not ten_hang:
                continue
            ten_hang = str(ten_hang).strip()
            if not ten_hang or "Tổng" in ten_hang:
                continue
            if "CHIẾT KHẤU" in ten_hang.upper() or "TT XÁC NHẬN" in ten_hang.upper():
                continue

            ncc = str(row[col_ncc]).strip() if col_ncc < len(row) and row[col_ncc] else ""
            if not ncc or ncc.lower() == "none":
                ncc = ""

            ma_hang = str(row[col_ma_hang]).strip() if col_ma_hang < len(row) and row[col_ma_hang] else ""
            dvt     = str(row[col_dvt]).strip()     if col_dvt < len(row) and row[col_dvt]     else ""

            try:
                gia_cost = int(float(row[col_gia_cost])) if col_gia_cost < len(row) and row[col_gia_cost] else 0
            except (ValueError, TypeError):
                gia_cost = 0

            branch_qty = {}
            for branch, col in branches.items():
                if col < len(row):
                    try:
                        qty = float(row[col] or 0)
                        if qty > 0:
                            branch_qty[branch] = int(qty)
                    except (ValueError, TypeError):
                        pass

            if branch_qty:
                products.append({
                    "ncc":        ncc,
                    "ma_hang":    ma_hang,
                    "ten_hang":   ten_hang,
                    "dvt":        dvt,
                    "gia_cost":   gia_cost,
                    "branch_qty": branch_qty,
                })

        all_data[sheet_name] = {
            "branches": list(branches.keys()),
            "products": products,
        }

    return all_data

# TAB 1 Tạo một phiếu XUẤT KHO ───────────────────────────────────────────────────────
def build_phieu(branch_name, items, now, logo_bytes=None, ghi_chu="Upload lên Link HPC (Không Kiot)", show_gia=True, index=0):
    doc = Document()

    # Căn lề trang A4
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(2)
    section.right_margin  = Cm(1.5)
    sectPr = section._sectPr
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:start"), "1")
    sectPr.append(pgNumType)

    # ── Header + Footer ──────────────────────────────────────────────────────

    def add_para(text, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.LEFT,
                 color=None, space_before=0, space_after=4):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = "Arial"
        if color:
            run.font.color.rgb = RGBColor(*color)
        return p

    # Tiêu đề
    ncc = items[0]["ncc"].upper() if items else "Không có NCC"

    p_ncc = doc.add_paragraph()
    p_ncc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ncc.paragraph_format.space_before = Pt(0)
    p_ncc.paragraph_format.space_after  = Pt(2)
    r_ncc = p_ncc.add_run(f"NHÀ CUNG CẤP: {ncc}")
    r_ncc.bold = True
    r_ncc.font.size = Pt(20)
    r_ncc.font.name = "Arial"
    r_ncc.font.color.rgb = RGBColor(220, 0, 0)

    p_gc = doc.add_paragraph()
    p_gc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_gc.paragraph_format.space_before = Pt(0)
    p_gc.paragraph_format.space_after  = Pt(6)
    r_gc = p_gc.add_run(f"⚠️ {ghi_chu}")
    r_gc.bold = True
    r_gc.italic = True
    r_gc.font.size = Pt(13)
    r_gc.font.name = "Arial"
    r_gc.font.color.rgb = RGBColor(192, 0, 0)

    add_para("PHIẾU XUẤT KHO", bold=True, size=18,
             align=WD_ALIGN_PARAGRAPH.CENTER,
             color=(43, 84, 126), space_before=20, space_after=2)

    # Đường kẻ dưới tiêu đề
    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_after = Pt(8)
    pPr = p_line._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:color"), "2B547E")
    pBdr.append(bottom)
    pPr.append(pBdr)

    ma_phieu = make_ma_phieu(now, index)
    ngay_tao = now.strftime("%d/%m/%Y")

    # Thông tin phiếu
    def info_line(label, value, value_color=None, value_bold=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r1 = p.add_run(label)
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.name = "Arial"
        r2 = p.add_run(value)
        r2.font.size = Pt(11)
        r2.font.name = "Arial"
        if value_color:
            r2.font.color.rgb = RGBColor(*value_color)
        if value_bold:
            r2.bold = True

    info_line("Mã phiếu: ", f"{ma_phieu}")
    info_line("Ngày tạo phiếu: ", f"{ngay_tao}")
    info_line("Trung Tâm nhận: ", branch_name, (43, 84, 126), value_bold=True)

    # Bảng hàng hóa
    col_widths = [Cm(0.8), Cm(2.2), Cm(5.5), Cm(1.3), Cm(1.8), Cm(2.2), Cm(2.5)]
    headers    = ["STT", "Mã hàng", "Tên hàng", "ĐVT", "Số lượng", "Đơn giá", "Thành tiền"]

    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"

    hdr = table.rows[0]
    for i, (cell, text, w) in enumerate(zip(hdr.cells, headers, col_widths)):
        cell.width = w
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, "FFFFFF")
        set_cell_border(cell, "000000")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0, 0, 0)

    total_qty    = 0
    total_amount = 0

    for idx, item in enumerate(items):
        qty    = item["qty"]
        amount = qty * item["gia_cost"]
        total_qty    += qty
        total_amount += amount
        shade = "FFFFFF"

        row = table.add_row()
        values = [
            str(idx + 1),
            item.get("ma_hang", ""),
            item.get("ten_hang", ""),
            item.get("dvt", ""),
            fmt_number(qty),
            fmt_number(item["gia_cost"]) if show_gia else "",
            fmt_number(amount) if show_gia else "",
        ]
        aligns = [
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.RIGHT,
            WD_ALIGN_PARAGRAPH.RIGHT,
        ]
        for cell, text, w, align in zip(row.cells, values, col_widths, aligns):
            cell.width = w
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_bg(cell, shade)
            set_cell_border(cell, "000000")
            p = cell.paragraphs[0]
            p.alignment = align
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.font.name = "Arial"

    # Dòng tổng cộng
    row_total = table.add_row()
    cells = row_total.cells

    cells[0].merge(cells[1]).merge(cells[2]).merge(cells[3])
    set_cell_bg(cells[0], "EEEEEE")
    set_cell_border(cells[0], "000000")
    p = cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run("TỔNG CỘNG")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0, 0, 0)

    for cell, text in [(cells[4], fmt_number(total_qty)),
                   (cells[5], ""),
                   (cells[6], fmt_number(total_amount) if show_gia else "")]:
        set_cell_bg(cell, "EEEEEE")
        set_cell_border(cell, "000000")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0, 0, 0)

    # ── Ký tên: chỉ Người Gửi và Người Nhận ────────────────────────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    sig_table = doc.add_table(rows=1, cols=2)
    sig_labels = ["Người Gửi", "Người Nhận"]
    for cell, label in zip(sig_table.rows[0].cells, sig_labels):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            tag = OxmlElement(f"w:{edge}")
            tag.set(qn("w:val"), "none")
            tcBorders.append(tag)
        tcPr.append(tcBorders)

        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(label)
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.name = "Arial"

        p2 = cell.add_paragraph("(Ký, ghi rõ họ tên)")
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.runs[0].italic = True
        p2.runs[0].font.size = Pt(9)
        p2.runs[0].font.name = "Arial"
        p2.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    p_break = doc.add_paragraph() 
    p_break.paragraph_format.space_before = Pt(0)
    p_break.paragraph_format.space_after  = Pt(0)
    run_break = p_break.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run_break._r.append(br)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

# TAB 2 Tạo phiếu chuyển hàng ───────────────────────────────────────────────────
def build_phieu_chuyen(ma_phieu, ngay_chuyen, chi_nhanh_chuyen, chi_nhanh_nhan,
                       ghi_chu_chuyen, ghi_chu_nhan, items, logo_bytes=None):
    doc = Document()

    section = doc.sections[0]
    section.page_width    = Cm(21)
    section.page_height   = Cm(29.7)
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(2)
    section.right_margin  = Cm(1.5)
    # ── Logo góc trên trái ──────────────────────────────────────────────────
    if logo_bytes:
        try:
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_logo.paragraph_format.space_before = Pt(0)
            p_logo.paragraph_format.space_after  = Pt(6)
            run_logo = p_logo.add_run()
            run_logo.add_picture(io.BytesIO(logo_bytes), height=Cm(1.2))
        except Exception:
            pass
    # ────────────────────────────────────────────────────────────────────────
    def add_para(text, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.LEFT,
                 color=None, space_before=0, space_after=4):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = "Arial"
        if color:
            run.font.color.rgb = RGBColor(*color)
        return p

    # Tiêu đề
    add_para("PHIẾU CHUYỂN HÀNG", bold=True, size=18,
             align=WD_ALIGN_PARAGRAPH.CENTER,
             color=(43, 84, 126), space_after=2)

    # Đường kẻ dưới tiêu đề
    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_after = Pt(8)
    pPr = p_line._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:color"), "2B547E")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Thông tin phiếu
    def info_line(label, value, value_color=None, value_bold=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r1 = p.add_run(label)
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.name = "Arial"
        r2 = p.add_run(value)
        r2.font.size = Pt(11)
        r2.font.name = "Arial"
        if value_color:
            r2.font.color.rgb = RGBColor(*value_color)
        if value_bold:
            r2.bold = True

    info_line("Mã phiếu: ",          ma_phieu)
    info_line("Ngày chuyển: ",        ngay_chuyen)
    info_line("Trung tâm chuyển: ",   chi_nhanh_chuyen, (43, 84, 126), value_bold=True)
    info_line("Trung tâm nhận: ",     chi_nhanh_nhan,   (43, 84, 126), value_bold=True)

    # Ghi chú bên chuyển
    if ghi_chu_chuyen:
        p_note = doc.add_paragraph()
        p_note.paragraph_format.space_before = Pt(2)
        p_note.paragraph_format.space_after  = Pt(2)
        r1 = p_note.add_run("Ghi chú bên chuyển: ")
        r1.bold = True; r1.font.size = Pt(11); r1.font.name = "Arial"
        r2 = p_note.add_run(ghi_chu_chuyen)
        r2.italic = True; r2.font.size = Pt(11); r2.font.name = "Arial"
        r2.font.color.rgb = RGBColor(192, 0, 0)

    # Ghi chú bên nhận
    if ghi_chu_nhan:
        p_note2 = doc.add_paragraph()
        p_note2.paragraph_format.space_before = Pt(2)
        p_note2.paragraph_format.space_after  = Pt(10)
        r1 = p_note2.add_run("Ghi chú bên nhận: ")
        r1.bold = True; r1.font.size = Pt(11); r1.font.name = "Arial"
        r2 = p_note2.add_run(ghi_chu_nhan)
        r2.italic = True; r2.font.size = Pt(11); r2.font.name = "Arial"
        r2.font.color.rgb = RGBColor(192, 0, 0)

    # Bảng hàng hóa
    col_widths = [Cm(1.0), Cm(2.5), Cm(6.0), Cm(2.0), Cm(2.5)]
    headers    = ["STT", "Mã hàng", "Tên hàng", "ĐVT", "SL chuyển"]

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    hdr = table.rows[0]
    for cell, text, w in zip(hdr.cells, headers, col_widths):
        cell.width = w
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, "2B547E")
        set_cell_border(cell, "2B547E")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(255, 255, 255)

    total_chuyen = 0
    total_nhan   = 0

    for idx, item in enumerate(items):
        sl_chuyen = item.get("sl_chuyen", 0)
        sl_nhan   = item.get("sl_nhan", sl_chuyen)  # mặc định = sl_chuyen nếu không có
        total_chuyen += sl_chuyen
        total_nhan   += sl_nhan

        row = table.add_row()
        values = [
                str(idx + 1),
                item.get("ma_hang", ""),
                item.get("ten_hang", ""),
                item.get("dvt", ""),
                fmt_number(sl_chuyen),
            ]
        aligns = [
                WD_ALIGN_PARAGRAPH.CENTER,
                WD_ALIGN_PARAGRAPH.LEFT,
                WD_ALIGN_PARAGRAPH.LEFT,
                WD_ALIGN_PARAGRAPH.CENTER,
                WD_ALIGN_PARAGRAPH.CENTER,
            ]
        for cell, text, w, align in zip(row.cells, values, col_widths, aligns):
            cell.width = w
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_bg(cell, "FFFFFF")
            set_cell_border(cell, "CCCCCC")
            p = cell.paragraphs[0]
            p.alignment = align
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.font.name = "Arial"

    # Dòng tổng cộng
    row_total = table.add_row()
    cells = row_total.cells
    cells[0].merge(cells[1]).merge(cells[2]).merge(cells[3])
    set_cell_bg(cells[0], "EFF4FA")
    set_cell_border(cells[0], "2B547E")
    p = cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run("TỔNG CỘNG")
    run.bold = True; run.font.size = Pt(10); run.font.name = "Arial"
    run.font.color.rgb = RGBColor(43, 84, 126)

    set_cell_bg(cells[4], "2B547E")
    set_cell_border(cells[4], "2B547E")
    p = cells[4].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(fmt_number(total_chuyen))
    run.bold = True; run.font.size = Pt(10); run.font.name = "Arial"
    run.font.color.rgb = RGBColor(255, 255, 255)

    # Ký tên
    doc.add_paragraph().paragraph_format.space_after = Pt(16)
    sig_table = doc.add_table(rows=1, cols=2)
    for cell, label in zip(sig_table.rows[0].cells, ["Bên Chuyển", "Bên Nhận"]):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            tag = OxmlElement(f"w:{edge}")
            tag.set(qn("w:val"), "none")
            tcBorders.append(tag)
        tcPr.append(tcBorders)
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(label)
        r1.bold = True; r1.font.size = Pt(10); r1.font.name = "Arial"
        p2 = cell.add_paragraph("(Ký, ghi rõ họ tên)")
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.runs[0].italic = True
        p2.runs[0].font.size = Pt(9)
        p2.runs[0].font.name = "Arial"
        p2.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # Page break để tách phiếu khi ghép
    p_break = doc.add_paragraph()
    p_break.paragraph_format.space_before = Pt(0)
    p_break.paragraph_format.space_after  = Pt(0)
    run_break = p_break.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run_break._r.append(br)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
# ─── Helpers dùng chung cho parse_excel_raw ──────────────────────────────────
SKIP_TT = {"tổng order", "tổng", "none", ""}

def find_col(keywords, search_rows):
    """Tìm index cột đầu tiên khớp với bất kỳ keyword nào trong các hàng cho trước."""
    for row in search_rows:
        for i, v in enumerate(row):
            if v and any(kw.lower() in str(v).lower() for kw in keywords):
                return i
    return None

# ─── Chuyển đổi file sai định dạng ───────────────────────────────────────────
def parse_excel_raw(file_bytes):
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    all_branches = {}

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if len(rows) < 4:
            continue

        row1 = rows[0]
        row2 = rows[1]
        row3 = rows[2]

        # Auto-detect: nếu hàng 1 chỉ có 1-2 ô (tiêu đề chung) → lệch 1 hàng
        row1_non_none = [v for v in row1 if v is not None]
        data_offset = 3
        is_offset_format = False
        if len(row1_non_none) <= 2:
            row2 = rows[1]   # TT
            row3 = rows[2]   # sub-header "Số đặt"
            data_offset = 4
            is_offset_format = True

        ...

        # Phát hiện định dạng:
        # - Định dạng mới (Kids): row2 có "Số đặt", tên TT ở row1
        # - Định dạng cũ (bóng kho): row3 có "Số đặt", tên TT ở row2
        has_so_dat_in_row2 = any(
            str(v).strip().lower() == "số đặt" for v in row2 if v
        )
        has_so_dat_in_row3 = any(
            str(v).strip().lower() == "số đặt" for v in row3 if v
        )

        if has_so_dat_in_row2 and not has_so_dat_in_row3:
            # ── ĐỊNH DẠNG MỚI: Hàng 1 = tên trung tâm, Hàng 2 = sub-header ──
            COL_NCC      = find_col(["nhà cung cấp", "ncc"],          [row2])
            COL_MA_HANG  = find_col(["mã hàng", "ma hang", "mã sp"],  [row2])
            COL_TEN_HANG = find_col(["tên hàng"],                      [row2])
            COL_DVT      = find_col(["đvt", "đơn vị tính"],            [row2])
            COL_GIA_COST = find_col(["đơn giá", "giá cost"],           [row2])

            if COL_NCC      is None: COL_NCC      = 2
            if COL_MA_HANG  is None: COL_MA_HANG  = 1
            if COL_TEN_HANG is None: COL_TEN_HANG = 3
            if COL_DVT      is None: COL_DVT      = 5
            if COL_GIA_COST is None: COL_GIA_COST = 7

            so_dat_cols = [i for i, v in enumerate(row2) if v and str(v).strip().lower() == "số đặt"]

            branch_list_final = []
            for i, val in enumerate(row1):
                if not val or str(val).strip().lower() in SKIP_TT:
                    continue
                ten_tt = str(val).strip()
                col_so_dat = next((sc for sc in so_dat_cols if sc >= i), None)
                if col_so_dat is None:
                    continue
                unique_name = ten_tt
                count = sum(1 for n, _ in branch_list_final if n == ten_tt or n.startswith(ten_tt + "_"))
                if count > 0:
                    unique_name = f"{ten_tt}_{count + 1}"
                branch_list_final.append((unique_name, col_so_dat))

            data_start = 3  # bỏ hàng 1+2 (header) và hàng 3 (Tổng Cộng)

            for branch_name, col_so_dat in branch_list_final:
                items = []
                for row in rows[data_start:]:
                    if not row or len(row) <= COL_TEN_HANG:
                        continue
                    ten_hang = row[COL_TEN_HANG]
                    if not ten_hang:
                        continue
                    ten_hang = str(ten_hang).strip()
                    if not ten_hang or "tổng" in ten_hang.lower():
                        continue

                    so_dat = 0
                    if col_so_dat < len(row) and row[col_so_dat]:
                        try:
                            so_dat = float(row[col_so_dat])
                        except (ValueError, TypeError):
                            so_dat = 0
                    if so_dat <= 0:
                        continue

                    try:
                        gia_cost = float(row[COL_GIA_COST]) if COL_GIA_COST < len(row) and row[COL_GIA_COST] else 0
                    except (ValueError, TypeError):
                        gia_cost = 0

                    items.append({
                        "NCC":       str(row[COL_NCC]).strip()      if COL_NCC < len(row) and row[COL_NCC] else "",
                        "Mã hàng":   str(row[COL_MA_HANG]).strip()  if COL_MA_HANG < len(row) and row[COL_MA_HANG] else "",
                        "Tên Hàng":  ten_hang,
                        "Cách đóng": str(row[COL_DVT]).strip()      if COL_DVT < len(row) and row[COL_DVT]  else "",
                        "Giá cost":  gia_cost,
                        "Số đặt":    int(so_dat),
                    })

                if items:
                    all_branches[branch_name] = items

        else:
            # ── ĐỊNH DẠNG CŨ: row2 = tên trung tâm, row3 = "Số đặt" ──
            # (áp dụng cho file bóng kho có tiêu đề ở dòng 1)
            COL_NCC      = find_col(["NCC"],                                                [row3, row2])
            COL_MA_HANG  = find_col(["mã hàng", "ma hang", "mã sp"],                       [row3, row2])
            COL_TEN_HANG = find_col(["Tên hàng"],                                           [row3, row2])
            COL_DVT      = find_col(["Đơn vị đặt hàng", "đvt", "đơn vị tính"],             [row3, row2])
            COL_GIA_COST = find_col(["Đơn giá/ Cái", "Đơn giá/Cái", "đơn giá", "giá cost"],[row3, row2])

            if COL_NCC      is None: COL_NCC      = 3
            if COL_MA_HANG  is None: COL_MA_HANG  = 1
            if COL_TEN_HANG is None: COL_TEN_HANG = 5
            if COL_DVT      is None: COL_DVT      = 9
            if COL_GIA_COST is None: COL_GIA_COST = 8

            branch_list_final = []
            for i in range(0, len(row2)):
                val = row2[i]
                if not val or not isinstance(val, str):
                    continue
                val = val.strip()
                if not val or val.lower() in SKIP_TT or "tháng" in val.lower():
                    continue
                if any(kw in val.lower() for kw in ["nhóm", "mã hàng", "phụ trách", "ncc", "vat", "tên hàng", "ảnh", "quy cách", "tổng order"]):
                    continue
                try:
                    float(val)
                    continue
                except ValueError:
                    pass
                col_so_dat = None
                for j in range(i, min(i + 10, len(row3))):
                    if row3[j] and str(row3[j]).strip() == "Số đặt":
                        col_so_dat = j
                        break
                if col_so_dat is None:
                    continue
                unique_name = val
                count = sum(1 for n, _ in branch_list_final if n == val or n.startswith(val + "_"))
                if count > 0:
                    unique_name = f"{val}_{count + 1}"
                branch_list_final.append((unique_name, col_so_dat))

            for branch_name, col_so_dat in branch_list_final:
                items = []
                for row in rows[data_offset:]:
                    if not row or len(row) <= COL_TEN_HANG:
                        continue
                    ten_hang = row[COL_TEN_HANG]
                    if not ten_hang:
                        continue
                    ten_hang = str(ten_hang).strip()
                    if not ten_hang or "tổng" in ten_hang.lower():
                        continue

                    so_dat = 0
                    if col_so_dat < len(row) and row[col_so_dat]:
                        try:
                            so_dat = float(row[col_so_dat])
                        except (ValueError, TypeError):
                            so_dat = 0
                    if so_dat <= 0:
                        continue

                    try:
                        gia_cost = float(row[COL_GIA_COST]) if COL_GIA_COST < len(row) and row[COL_GIA_COST] else 0
                    except (ValueError, TypeError):
                        gia_cost = 0

                    items.append({
                        "NCC":       str(row[COL_NCC]).strip()      if COL_NCC < len(row) and row[COL_NCC] else "",
                        "Mã hàng":   str(row[COL_MA_HANG]).strip()  if COL_MA_HANG < len(row) and row[COL_MA_HANG] else "",
                        "Tên Hàng":  ten_hang,
                        "Cách đóng": str(row[COL_DVT]).strip()      if COL_DVT < len(row) and row[COL_DVT]  else "",
                        "Giá cost":  gia_cost,
                        "Số đặt":    so_dat,
                    })

                if items:
                    all_branches[branch_name] = items

    return all_branches
def export_to_standard_excel(all_branches: dict) -> bytes:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Dữ liệu chuẩn"

    # Gom tất cả sản phẩm unique (NCC + Tên Hàng + ĐVT + Đơn giá)
    # Key: (NCC, Tên Hàng, ĐVT, Đơn giá) — Value: {tên_trung_tâm: số_đặt}
    product_map = {}
    branch_names = list(all_branches.keys())

    for branch_name, items in all_branches.items():
        for item in items:
            key = (item["NCC"], item["Tên Hàng"], item["Cách đóng"], item["Giá cost"])
            if key not in product_map:
                product_map[key] = {}
            product_map[key][branch_name] = item["Số đặt"]

    # Header: NCC | Tên hàng | ĐVT | Đơn giá | TT1 | TT2 | ...
    headers = ["NCC", "Tên hàng", "ĐVT", "Đơn giá"] + branch_names
    ws.append(headers)

    # Dữ liệu
    for (ncc, ten_hang, dvt, don_gia), branch_qty in product_map.items():
        row = [ncc, ten_hang, dvt, don_gia]
        for branch in branch_names:
            row.append(branch_qty.get(branch, 0))
        ws.append(row)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()
# ─── Tạo file Excel mẫu ──────────────────────────────────────────────────────
def create_sample_excel():
    wb = openpyxl.Workbook()

    # ── Sheet 1: Dạng thường (không có cột NCC) ──────────────────────────────
    ws1 = wb.active
    ws1.title = "Sheet_Thuong"

    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    header_font  = Font(name="Arial", bold=True, color="FFFFFF", size=11)
    header_fill  = PatternFill("solid", fgColor="2B547E")
    center_align = Alignment(horizontal="center", vertical="center")
    thin_border  = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"),  bottom=Side(style="thin")
    )

    # Hàng 1: Tên chi nhánh (từ cột C trở đi)
    ws1["A1"] = "Tên hàng"
    ws1["B1"] = "Đơn giá (cost)"
    ws1["C1"] = "Trung Tâm A"
    ws1["D1"] = "Trung Tâm B"
    ws1["E1"] = "Trung Tâm C"

    for cell in ws1[1]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center_align

    # Hàng 2: Sub-header
    ws1["A2"] = "Tên hàng"
    ws1["B2"] = "Giá cost"
    ws1["C2"] = "Số đặt"
    ws1["D2"] = "Số đặt"
    ws1["E2"] = "Số đặt"

    # Hàng 3: Tiêu đề cột (bỏ trống hoặc ghi chú)
    ws1["A3"] = "(Bỏ trống hoặc ghi chú)"

    # Dữ liệu mẫu từ hàng 4
    sample_data = [
        ("Sữa tươi Vinamilk 1L", 28000, 10, 5, 8),
        ("Bánh mì sandwich", 15000, 20, 15, 12),
        ("Nước suối Lavie 500ml", 5000, 50, 30, 40),
        ("Snack khoai tây", 12000, 15, 10, 20),
    ]
    for row_data in sample_data:
        ws1.append(list(row_data))

    # Điều chỉnh độ rộng cột
    ws1.column_dimensions["A"].width = 30
    ws1.column_dimensions["B"].width = 18
    for col in ["C", "D", "E"]:
        ws1.column_dimensions[col].width = 14

    # ── Sheet 2: Dạng có NCC (như sheet Kids) ────────────────────────────────
    ws2 = wb.create_sheet("Sheet_CoNCC")

    ws2["A1"] = "NCC"
    ws2["B1"] = "Tên hàng"
    ws2["C1"] = "ĐVT"
    ws2["D1"] = "Đơn giá (cost)"
    ws2["E1"] = "Trung Tâm A"
    ws2["F1"] = "Trung Tâm B"
    ws2["G1"] = "Trung Tâm C"

    for cell in ws2[1]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center_align

    ws2["A2"] = "NCC"
    ws2["B2"] = "Tên hàng"
    ws2["C2"] = "ĐVT"
    ws2["D2"] = "Giá cost"
    ws2["E2"] = "Số đặt"
    ws2["F2"] = "Số đặt"
    ws2["G2"] = "Số đặt"

    ws2["A3"] = "(Bỏ trống)"

    sample_data2 = [
        ("Vinamilk",   "Sữa bột Dielac 900g",    "Hộp",   285000, 5, 3, 4),
        ("TH True",    "Sữa tươi TH 180ml",       "Hộp",   7500,  30, 20, 25),
        ("Nestle",     "Milo hộp 180ml",           "Hộp",   9500,  20, 15, 18),
        ("Abbott",     "Ensure Gold 850g",         "Lon",   450000, 2, 1, 3),
    ]
    for row_data in sample_data2:
        ws2.append(list(row_data))

    ws2.column_dimensions["A"].width = 14
    ws2.column_dimensions["B"].width = 28
    ws2.column_dimensions["C"].width = 10
    ws2.column_dimensions["D"].width = 18
    for col in ["E", "F", "G"]:
        ws2.column_dimensions[col].width = 14

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()

# ─── Giao diện chính ──────────────────────────────────────────────────────────

# ── Định dạng gốc ──────────────────────────────────────────────
with st.expander("📋 Mẫu excel", expanded=False):
        st.markdown("""
### 📌 Cấu trúc file Excel chuẩn

**Hàng 1:** Tên Trung Tâm (từ cột H trở đi, mỗi TT chiếm 2 cột, merge lại)

|  |  |  |  |  |  |  | Trung Tâm A | | Trung Tâm B | | ... |
|---|---|---|---|---|---|---|---|---|---|---|---|
| NCC | Mã hàng | Tên hàng | ĐVT | Đơn giá | Số đặt | Thành Tiền | Số đặt | Thành tiền | Số đặt | Thành tiền | ... |
| Duka | DK81397 | Xe SUV cảnh sát | Hộp | 285.000 | 8 | 2.280.000 | 5 | 1.425.000 | 3 | 855.000 | ... |

- Cột A–E: NCC, Mã hàng, Tên hàng, ĐVT, Đơn giá
- Cột F–G: Tổng Số đặt + Thành Tiền (không có tên TT)
- Cột H trở đi: mỗi Trung Tâm 2 cột (Số đặt + Thành tiền)
- Hàng 3: Tổng Cộng (tự động)
- Hàng 4+: Dữ liệu sản phẩm
        """)

        def create_sample_hpc():
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from openpyxl.utils import get_column_letter

            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Thanh toán Toystore"

            hdr_font = Font(name="Arial", bold=True, color="FFFFFF", size=10)
            hdr_fill = PatternFill("solid", fgColor="2B547E")
            tt_fill  = PatternFill("solid", fgColor="1F4E79")
            tot_fill = PatternFill("solid", fgColor="D6E4F0")
            center   = Alignment(horizontal="center", vertical="center", wrap_text=True)
            left_al  = Alignment(horizontal="left",   vertical="center")
            right_al = Alignment(horizontal="right",  vertical="center")
            thin     = Side(style="thin", color="CCCCCC")
            bdr      = Border(left=thin, right=thin, top=thin, bottom=thin)

            trung_tams = ["Trung Tâm A", "Trung Tâm B", "Trung Tâm C", "Trung Tâm D"]
            products = [
                ("Mai Loan", "B2109817", "Giỏ hoa quả xách tay",       "Bộ",  315000, [5, 0, 2, 1]),
                ("Mai Loan", "B2109818", "Khay bánh sữa thực phẩm",    "Bộ",  172000, [1, 0, 0, 1]),
                ("Mai Loan", "B2109819", "Set 4 khay thực phẩm",       "Bộ",  415000, [2, 1, 0, 0]),
                ("Duka",     "D2109820", "Xe SUV cảnh sát",             "Hộp", 285000, [3, 2, 1, 0]),
                ("Duka",     "D2109821", "Bộ đồ chơi bác sĩ nha khoa", "Bộ",  348000, [1, 1, 2, 1]),
            ]

            TT_START = 8

            # Hàng 1: Tên TT
            for i, tt in enumerate(trung_tams):
                col = TT_START + i * 2
                cell = ws.cell(1, col)
                cell.value = tt
                cell.font = Font(name="Arial", bold=True, color="FFFFFF", size=10)
                cell.fill = tt_fill
                cell.alignment = center
                ws.merge_cells(start_row=1, start_column=col, end_row=1, end_column=col + 1)
            ws.row_dimensions[1].height = 30

            # Hàng 2: Sub-header
            row2_vals = ["NCC", "Mã hàng", "Tên hàng", "ĐVT", "Đơn giá", "Số đặt", "Thành Tiền"]
            for _ in trung_tams:
                row2_vals += ["Số đặt", "Thành tiền"]
            for c, val in enumerate(row2_vals, 1):
                cell = ws.cell(2, c)
                cell.value = val
                cell.font = hdr_font
                cell.fill = hdr_fill
                cell.alignment = center
                cell.border = bdr
            ws.row_dimensions[2].height = 25

            # Hàng 3: Tổng Cộng
            total_by_tt  = [0] * len(trung_tams)
            total_amt_tt = [0] * len(trung_tams)
            grand_qty = grand_amt = 0
            for _, _, _, _, gia, qtys in products:
                for i, q in enumerate(qtys):
                    total_by_tt[i]  += q
                    total_amt_tt[i] += q * gia
                    grand_qty += q
                    grand_amt += q * gia

            ws.cell(3, 3).value = "Tổng Cộng"
            ws.cell(3, 3).font = Font(name="Arial", bold=True, size=10, color="000000")
            ws.cell(3, 4).value = 0
            for c, val in [(6, grand_qty), (7, grand_amt)]:
                cell = ws.cell(3, c)
                cell.value = val
                cell.font = Font(name="Arial", bold=True, size=10, color="000000")
                cell.fill = tot_fill; cell.alignment = center; cell.border = bdr
            for i in range(len(trung_tams)):
                col = TT_START + i * 2
                for c, val in [(col, total_by_tt[i]), (col + 1, total_amt_tt[i])]:
                    cell = ws.cell(3, c)
                    cell.value = val
                    cell.font = Font(name="Arial", bold=True, size=10, color="000000")
                    cell.fill = tot_fill; cell.alignment = center; cell.border = bdr

            # Hàng 4+: Dữ liệu
            for row_idx, (ncc, ma, ten, dvt, gia, qtys) in enumerate(products, 4):
                shade = "F5F9FF" if row_idx % 2 == 0 else "FFFFFF"
                rfill = PatternFill("solid", fgColor=shade)
                p_qty = sum(qtys); p_amt = p_qty * gia
                for c, val, aln in [(1,ncc,center),(2,ma,left_al),(3,ten,left_al),(4,dvt,center),(5,gia,right_al),(6,p_qty,center),(7,p_amt,right_al)]:
                    cell = ws.cell(row_idx, c)
                    cell.value = val
                    cell.font = Font(name="Arial", size=10, color="000000")
                    cell.fill = rfill; cell.alignment = aln; cell.border = bdr
                for i, qty in enumerate(qtys):
                    col = TT_START + i * 2
                    amt = qty * gia if qty else 0
                    for c, val in [(col, qty or None), (col + 1, amt or None)]:
                        cell = ws.cell(row_idx, c)
                        cell.value = val
                        cell.font = Font(name="Arial", size=10, color="000000")
                        cell.fill = rfill; cell.alignment = center; cell.border = bdr

            # Độ rộng cột
            ws.column_dimensions["A"].width = 12
            ws.column_dimensions["B"].width = 12
            ws.column_dimensions["C"].width = 35
            ws.column_dimensions["D"].width = 8
            ws.column_dimensions["E"].width = 12
            ws.column_dimensions["F"].width = 10
            ws.column_dimensions["G"].width = 14
            for i in range(len(trung_tams) * 2):
                ws.column_dimensions[get_column_letter(TT_START + i)].width = 11
            ws.freeze_panes = "F3"

            buf = io.BytesIO()
            wb.save(buf)
            buf.seek(0)
            return buf.getvalue()

        st.download_button(
            label="⬇️ Tải file Excel mẫu",
            data=create_sample_hpc(),
            file_name="Mau_File.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )


if page == "xuat_kho":
    st.subheader("🏭 Tạo phiếu xuất kho")
# ── Upload file ──────────────────────────────────────────────────────────────
    st.markdown("<div style='margin-bottom:-20px'><b>Chọn loại ghi chú cho phiếu:</b> <span style='color:red'>*</span></div>", unsafe_allow_html=True)
    ghi_chu_option = st.selectbox(
    "",
    options=[
    "--- Chọn loại ghi chú ---",
    "Upload lên Link HPC (Nhập hàng KHÔNG Kiot)",
    "Upload lên Link nhập hàng (Nhập hàng CÓ Kiot)"
    ]
    )
    if ghi_chu_option == "--- Chọn loại ghi chú ---":
        st.warning("⚠️ Vui lòng chọn loại ghi chú trước khi tạo phiếu!")
    else:
        uploaded = st.file_uploader(
            "Kéo thả hoặc chọn file Excel",
            type=["xlsx"],
            help="File Excel có các sheet chứa danh sách hàng theo trung tâm"
        )

        # Đọc logo một lần
        logo_bytes = load_logo_from_url(LOGO_URL)

        if uploaded:
            excel_basename = os.path.splitext(uploaded.name)[0]
            st.success(f"Đã tải file: **{uploaded.name}**")

            with st.spinner("Đang đọc file Excel..."):
                try:
                    all_data = parse_excel(uploaded.read())
                    st.session_state["data_xuat_kho"] = all_data
                except Exception as e:
                    st.error(f"Lỗi đọc file: {e}")
                    st.stop()

            total_branches = sum(
                len([b for b in data["branches"]
                    if any(b in p["branch_qty"] for p in data["products"])])
                for data in all_data.values()
            )
            total_sheets = len(all_data)

            col1, col2 = st.columns(2)
            col1.metric("Số sheet", total_sheets)
            col2.metric("Tổng số Trung Tâm có hàng", total_branches)

            st.divider()

            sheet_names = list(all_data.keys())
            selected_sheets = st.multiselect(
                "Chọn sheet cần tạo phiếu",
                options=sheet_names,
                default=sheet_names,
            )

            # ── Bảng kê tổng hợp — chỉ hiện sheet đã chọn ────────────────
            if selected_sheets:
                import pandas as pd

                st.markdown("#### 📋 Bảng kê tổng hợp")

                for sheet_name, data in all_data.items():
                    if sheet_name not in selected_sheets:  # chỉ hiện sheet đã chọn
                        continue
                    products  = data["products"]
                    branches  = data["branches"]

                    tt_co_hang = [b for b in branches if any(b in p["branch_qty"] for p in products)]
                    if not tt_co_hang:
                        continue

                    ncc_map = {}
                    for p in products:
                        ncc_map.setdefault(p["ncc"], []).append(p)

                    for ncc_name, items_ncc in ncc_map.items():
                        rows = []
                        tong_tt = {tt: 0 for tt in tt_co_hang}
                        tong_tien_ncc = 0

                        for p in items_ncc:
                            tong_sp = 0
                            tt_qty = {}
                            for tt in tt_co_hang:
                                qty = p["branch_qty"].get(tt, 0)
                                tt_qty[tt] = qty if qty else ""
                                tong_tt[tt] += qty
                                tong_sp += qty
                            row = {
                                "Mã hàng":    p.get("ma_hang", ""),
                                "Tên hàng":   p["ten_hang"],
                                "ĐVT":        p.get("dvt", ""),
                                "Đơn giá":    f"{int(p['gia_cost']):,}",
                                "Tổng SL":    tong_sp,
                                "Thành tiền": f"{int(tong_sp * p['gia_cost']):,}",
                                **tt_qty,
                            }
                            tong_tien_ncc += tong_sp * p["gia_cost"]
                            rows.append(row)

                        tong_row = {
                            "Mã hàng": "", "Tên hàng": "TỔNG CỘNG",
                            "ĐVT": "", "Đơn giá": "",
                        }
                        tong_sl = 0
                        for tt in tt_co_hang:
                            tong_row[tt] = tong_tt[tt] if tong_tt[tt] else ""
                            tong_sl += tong_tt[tt]
                        tong_row["Tổng SL"]    = tong_sl
                        tong_row["Thành tiền"] = f"{int(tong_tien_ncc):,}"
                        rows.append(tong_row)

                        tong_tien_str = f"{int(tong_tien_ncc):,} đ"
                        with st.expander(
                            f"🏭 [{sheet_name}] {ncc_name} — {len(items_ncc)} sản phẩm | "
                            f"{len(tt_co_hang)} Trung Tâm | Tổng tiền: {tong_tien_str}",
                            expanded=False
                        ):
                            st.dataframe(
                                pd.DataFrame(rows),
                                use_container_width=True,
                                hide_index=True,
                            )

            st.divider()
            show_gia = st.checkbox("Hiển thị đơn giá trong phiếu", value=True)
            selected_date = st.date_input("📅 Ngày tạo phiếu", value=datetime.now().date())
            if st.button("🖨️ Tạo phiếu xuất kho", type="primary", use_container_width=True):
                from docxcompose.composer import Composer
                now = datetime.combine(selected_date, datetime.now().time())
                count = 0

                progress = st.progress(0, text="Đang tạo phiếu...")

                tasks = []
                for sheet_name in selected_sheets:
                    data = all_data.get(sheet_name, {})
                    for branch in data.get("branches", []):
                        items = [
                            {**p, "qty": p["branch_qty"][branch]}
                            for p in data["products"]
                            if branch in p["branch_qty"]
                        ]
                        if items:
                            tasks.append((sheet_name, branch, items))

                # ── Tạo trang bảng kê đầu (không có header/footer) ───────
                summary_doc = Document()
                sec = summary_doc.sections[0]
                sec.page_width    = Cm(21)
                sec.page_height   = Cm(29.7)
                sec.top_margin    = Cm(1.5)
                sec.bottom_margin = Cm(1.5)
                sec.left_margin   = Cm(2)
                sec.right_margin  = Cm(1.5)

                # Tắt header/footer trang bảng kê
                sec.header.is_linked_to_previous = False
                sec.footer.is_linked_to_previous = False
                for p in sec.header.paragraphs:
                    p.clear()
                for p in sec.footer.paragraphs:
                    p.clear()

                # Tiêu đề bảng kê
                p_title = summary_doc.add_paragraph()
                p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_title.paragraph_format.space_before = Pt(20)
                p_title.paragraph_format.space_after  = Pt(4)
                r_title = p_title.add_run("BẢNG KÊ TỔNG HỢP PHIẾU XUẤT KHO")
                r_title.bold = True
                r_title.font.size = Pt(16)
                r_title.font.name = "Arial"
                r_title.font.color.rgb = RGBColor(43, 84, 126)

                p_date = summary_doc.add_paragraph()
                p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_date.paragraph_format.space_after = Pt(16)
                r_date = p_date.add_run(f"Ngày: {now.strftime('%d/%m/%Y')}")
                r_date.font.size = Pt(11)
                r_date.font.name = "Arial"

                # Bảng kê
                col_w = [Cm(1.2), Cm(5.5), Cm(5.5), Cm(4)]
                tbl = summary_doc.add_table(rows=1, cols=4)
                tbl.style = "Table Grid"
                for cell, text, w in zip(tbl.rows[0].cells,
                                         ["STT", "Trung Tâm", "Mã phiếu", "Số lượng SP"],
                                         col_w):
                    cell.width = w
                    set_cell_bg(cell, "FFFFFF")
                    set_cell_border(cell, "000000")
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after  = Pt(3)
                    run = p.add_run(text)
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.name = "Arial"

                grand_total_sp = 0
                for i, (sheet_name, branch, items) in enumerate(tasks):
                    ma_phieu_i  = make_ma_phieu(now, i)
                    tong_sp_i   = sum(it["qty"] for it in items)
                    grand_total_sp += tong_sp_i
                    shade = "F0F4F8" if i % 2 == 0 else "FFFFFF"
                    row = tbl.add_row()
                    for cell, text, w, align in zip(
                        row.cells,
                        [str(i+1), branch, ma_phieu_i, fmt_number(tong_sp_i)],
                        col_w,
                        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                         WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER]
                    ):
                        cell.width = w
                        set_cell_bg(cell, shade)
                        set_cell_border(cell, "CCCCCC")
                        p = cell.paragraphs[0]
                        p.alignment = align
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.space_after  = Pt(2)
                        run = p.add_run(text)
                        run.font.size = Pt(10)
                        run.font.name = "Arial"

                # Dòng tổng
                row_total = tbl.add_row()
                for cell, text, w, align in zip(
                    row_total.cells,
                    ["", "TỔNG CỘNG", "", fmt_number(grand_total_sp)],
                    col_w,
                    [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                     WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER]
                ):
                    cell.width = w
                    set_cell_bg(cell, "EEEEEE")
                    set_cell_border(cell, "000000")
                    p = cell.paragraphs[0]
                    p.alignment = align
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after  = Pt(3)
                    run = p.add_run(text)
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.name = "Arial"

                # Thêm section break sau bảng kê để tách hoàn toàn
                p_sep = summary_doc.add_paragraph()
                p_sep.paragraph_format.space_before = Pt(0)
                p_sep.paragraph_format.space_after  = Pt(0)
                pPr = p_sep._p.get_or_add_pPr()
                sectPr_new = OxmlElement("w:sectPr")
                pgSz = OxmlElement("w:pgSz")
                pgSz.set(qn("w:w"), "12240")
                pgSz.set(qn("w:h"), "16838")
                sectPr_new.append(pgSz)
                pPr.append(sectPr_new)

                # ── Ghép bảng kê + các phiếu ─────────────────────────────
                base_doc = None
                composer = None

                for i, (sheet_name, branch, items) in enumerate(tasks):
                    docx_bytes = build_phieu(branch, items, now, logo_bytes, ghi_chu_option, show_gia, index=i)
                    doc = Document(io.BytesIO(docx_bytes))

                    if i == 0:
                        base_doc = doc
                        composer = Composer(base_doc)
                    else:
                        composer.append(doc)

                    count += 1
                    progress.progress((i + 1) / len(tasks),
                                    text=f"Đang tạo: {branch} ({i+1}/{len(tasks)})")

                # Ghép bảng kê vào đầu
                final_composer = Composer(summary_doc)
                final_composer.append(base_doc)

                progress.empty()

                buf = io.BytesIO()
                final_composer.save(buf)
                buf.seek(0)

                st.success(f"✅ Đã tạo xong **{count} phiếu xuất kho**!")
                st.download_button(
                    label=f"⬇️ Tải về {excel_basename}.docx",
                    data=buf,
                    file_name=f"{excel_basename}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )

# ─── Tab 2: Phiếu chuyển hàng ────────────────────────────────────────────────
elif page == "chuyen_hang":
    col_title, col_btn = st.columns([3, 1])
    with col_title:
        st.subheader("🚚 Tạo phiếu chuyển hàng")
        st.caption("Upload file Excel → Tự động tạo phiếu chuyển hàng Word cho từng trung tâm")
    with col_btn:
        st.markdown("""
        <a href="https://docs.google.com/spreadsheets/d/1M5HGPswsqUXFcv-9Pa5zg9u-xJkAqiq2NH88biuUBRs/edit?gid=785623052#gid=785623052" target="_blank" style="
            display: block;
            background: #2B547E;
            color: white;
            text-align: center;
            padding: 12px 16px;
            border-radius: 8px;
            font-weight: 800;
            font-size: 15px;
            text-decoration: none;
            box-shadow: 0 4px 12px rgba(231,76,60,0.4);
            letter-spacing: 0.3px;
        ">🔗 Link chuyển hàng</a>
        """, unsafe_allow_html=True)

    logo_bytes_tab3 = load_logo_from_url(LOGO_URL)

    GHI_CHU_CHUYEN_MAC_DINH = "YÊU CẦU TRUNG TÂM XÁC NHẬN HÀNG TRÊN KIOT"
    CHI_NHANH_CHUYEN_MAC_DINH = "Kho MH-VP"

    col_a, col_b = st.columns(2)
    with col_a:
        chi_nhanh_chuyen_inp = st.text_input("Người gửi", value=CHI_NHANH_CHUYEN_MAC_DINH)
        ghi_chu_chuyen_inp   = st.text_input("Ghi chú", value=GHI_CHU_CHUYEN_MAC_DINH)
    with col_b:
        ghi_chu_nhan_inp     = st.text_input("Người nhận", placeholder="Tuỳ chọn")
        ngay_chuyen_inp      = st.date_input("📅 Ngày chuyển", value=datetime.now().date())

    st.divider()

    # ── Phần 1: Nhập tay ──────────────────────────────────────────────────────
    st.caption("Mã hàng | Tên hàng | ĐVT | SL chuyển")

    if "chuyen_items" not in st.session_state:
        st.session_state.chuyen_items = [
            {"ma_hang": "", "ten_hang": "", "dvt": "", "sl_chuyen": 0, "gia_chuyen": 0}
        ]

    chi_nhanh_nhan_inp = st.text_input("Trung tâm nhận", placeholder="VD: KVRC", key="cn_nhan_tap")

    for i, item in enumerate(st.session_state.chuyen_items):
        c1, c2, c3, c4, c5, c6 = st.columns([2, 4, 1.5, 1.5, 2, 0.8])
        item["ma_hang"]    = c1.text_input("Mã hàng",     value=item["ma_hang"],    key=f"ma_{i}",  label_visibility="collapsed", placeholder="Mã hàng")
        item["ten_hang"]   = c2.text_input("Tên hàng",    value=item["ten_hang"],   key=f"ten_{i}", label_visibility="collapsed", placeholder="Tên hàng")
        item["dvt"]        = c3.text_input("ĐVT",         value=item.get("dvt",""), key=f"dvt_{i}", label_visibility="collapsed", placeholder="ĐVT")
        item["sl_chuyen"]  = c4.number_input("SL chuyển", value=item["sl_chuyen"],  key=f"slc_{i}", label_visibility="collapsed", min_value=0)
        if c6.button("🗑️", key=f"del_{i}") and len(st.session_state.chuyen_items) > 1:
            st.session_state.chuyen_items.pop(i)
            st.rerun()

    tong_sl = sum(it["sl_chuyen"] for it in st.session_state.chuyen_items)
    st.markdown(f"<div style='text-align:right; font-weight:700; color:#2B547E; font-size:14px; margin-top:4px;'>Tổng số lượng: {tong_sl:,}</div>", unsafe_allow_html=True)

    if st.button("➕ Thêm dòng", use_container_width=True):
        st.session_state.chuyen_items.append(
            {"ma_hang": "", "ten_hang": "", "dvt": "", "sl_chuyen": 0, "gia_chuyen": 0}
        )
        st.rerun()

    if st.button("🖨️ Tạo phiếu chuyển hàng", type="primary", use_container_width=True, key="btn_tap"):
        if not chi_nhanh_nhan_inp:
            st.warning("⚠️ Vui lòng nhập trung tâm nhận!")
        else:
            valid_items = [it for it in st.session_state.chuyen_items if it["ten_hang"].strip()]
            if not valid_items:
                st.warning("⚠️ Vui lòng nhập ít nhất 1 mặt hàng!")
            else:
                now_tap     = datetime.combine(ngay_chuyen_inp, datetime.now().time())
                ma_phieu_tap = f"PCH-{now_tap.strftime('%Y%m%d')}-{now_tap.strftime('%H%M%S')}"
                ngay_str_tap = now_tap.strftime("%d/%m/%Y %H:%M")

                docx_bytes_tap = build_phieu_chuyen(
                    ma_phieu         = ma_phieu_tap,
                    ngay_chuyen      = ngay_str_tap,
                    chi_nhanh_chuyen = chi_nhanh_chuyen_inp,
                    chi_nhanh_nhan   = chi_nhanh_nhan_inp,
                    ghi_chu_chuyen   = ghi_chu_chuyen_inp,
                    ghi_chu_nhan     = ghi_chu_nhan_inp,
                    items            = valid_items,
                    logo_bytes       = logo_bytes_tab3,
                )
                safe_cn  = chi_nhanh_nhan_inp.strip().replace("/", "-")
                st.session_state["last_phieu_tap"] = {
                    "bytes":       docx_bytes_tap,
                    "ma_phieu":    ma_phieu_tap,
                    "ngay_chuyen": ngay_str_tap,
                    "noi_chuyen":  chi_nhanh_chuyen_inp,
                    "tt_nhan":     chi_nhanh_nhan_inp,
                    "items":       valid_items,
                    "filename":    f"[TT-{safe_cn}] KHO TỔNG_{ma_phieu_tap}.docx",
                }

    # ── Nằm NGOÀI if button — giữ qua rerun ──
    if "last_phieu_tap" in st.session_state:
        p = st.session_state["last_phieu_tap"]
        st.success(f"✅ Đã tạo phiếu **{p['ma_phieu']}** thành công!")

        col_dl, col_yes, col_no = st.columns([3, 1.5, 1])

        col_dl.download_button(
            label     = f"⬇️ Tải phiếu",
            data      = p["bytes"],
            file_name = p["filename"],
            mime      = "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key       = "dl_phieu_tap",
        )
        col_yes.button("📋 Lưu lên Sheet", use_container_width=True, key="yes_sheet_tap")
        col_no.button("✖ Bỏ qua", use_container_width=True, key="no_sheet_tap")

        if st.session_state.get("yes_sheet_tap"):
            ok = log_phieu_chuyen_to_sheet(
                ma_phieu    = p["ma_phieu"],
                ngay_chuyen = p["ngay_chuyen"],
                noi_chuyen  = p["noi_chuyen"],
                tt_nhan     = p["tt_nhan"],
                items       = p["items"],
            )
            if ok:
                st.success("✅ Đã lưu lên Google Sheet thành công!")
            del st.session_state["last_phieu_tap"]

        if st.session_state.get("no_sheet_tap"):
            del st.session_state["last_phieu_tap"]
            st.rerun()
    # ── Phần 2: Upload file Excel ──────────────────────────────────────────────
    st.markdown("#### 📂 Tạo phiếu từ file Excel")
    uploaded_chuyen = st.file_uploader(
        "Kéo thả hoặc chọn file Excel order",
        type=["xlsx"],
        key="chuyen_uploader"
    )

    if uploaded_chuyen:
        chuyen_basename = os.path.splitext(uploaded_chuyen.name)[0]
        file_bytes_ch = uploaded_chuyen.read()
        st.session_state["chuyen_file_bytes"] = file_bytes_ch
        st.session_state["chuyen_basename"]   = chuyen_basename
    
    if st.session_state.get("chuyen_file_bytes"):
        file_bytes_ch   = st.session_state["chuyen_file_bytes"]
        chuyen_basename = st.session_state.get("chuyen_basename", "phieu_chuyen")

        # Đọc danh sách sheet
        wb_peek = openpyxl.load_workbook(io.BytesIO(file_bytes_ch), read_only=True)
        sheet_names_ch = wb_peek.sheetnames
        wb_peek.close()

        selected_sheets_ch = st.multiselect(
            "Chọn sheet cần tạo phiếu",
            options=sheet_names_ch,
            default=sheet_names_ch,
            key="ms_sheets_chuyen"
        )

        if selected_sheets_ch and st.button("📖 Đọc file", use_container_width=True, key="btn_read_chuyen"):
            with st.spinner("Đang đọc file Excel..."):
                try:
                    all_branches_chuyen = parse_excel_raw(file_bytes_ch)
                    st.session_state["data_chuyen_hang"] = all_branches_chuyen
                except Exception as e:
                    st.error(f"Lỗi đọc file: {e}")
                    st.stop()

        all_branches_chuyen = st.session_state.get("data_chuyen_hang", {})

        if all_branches_chuyen:
            import pandas as pd

            # Đảo ngược data: {TT: [items]} → {NCC: {ten_hang: {TT: qty, ...}}}
            ncc_map_ch = {}
            all_tt_ch  = list(all_branches_chuyen.keys())

            for branch_name, items in all_branches_chuyen.items():
                for it in items:
                    ncc  = it.get("NCC", "") or "Không rõ NCC"
                    key  = (it.get("Mã hàng",""), it.get("Tên Hàng",""), it.get("Cách đóng",""), it.get("Giá cost", 0))
                    if ncc not in ncc_map_ch:
                        ncc_map_ch[ncc] = {}
                    if key not in ncc_map_ch[ncc]:
                        ncc_map_ch[ncc][key] = {}
                    ncc_map_ch[ncc][key][branch_name] = ncc_map_ch[ncc][key].get(branch_name, 0) + it.get("Số đặt", 0)

            st.markdown("#### 📋 Bảng kê tổng hợp")

            for ncc_name, products_ch in ncc_map_ch.items():
                tt_co_hang = [
                    tt for tt in all_tt_ch
                    if any(tt in qty_map for qty_map in products_ch.values())
                ]
                rows_ch       = []
                tong_tt_ch    = {tt: 0 for tt in tt_co_hang}
                tong_tien_ncc = 0
                sp_count      = 0

                for (ma_hang, ten_hang, dvt, gia_cost), qty_map in products_ch.items():
                    tong_sp = sum(qty_map.get(tt, 0) for tt in tt_co_hang)
                    if tong_sp == 0:
                        continue
                    sp_count += 1
                    tt_qty = {tt: (qty_map.get(tt, 0) or "") for tt in tt_co_hang}
                    for tt in tt_co_hang:
                        tong_tt_ch[tt] += qty_map.get(tt, 0)
                    tong_tien = tong_sp * gia_cost
                    tong_tien_ncc += tong_tien
                    rows_ch.append({
                        "Mã hàng":    ma_hang,
                        "Tên hàng":   ten_hang,
                        "ĐVT":        dvt,
                        "Đơn giá":    f"{int(gia_cost):,}" if gia_cost else "",
                        "Tổng SL":    tong_sp,
                        "Thành tiền": f"{int(tong_tien):,}",
                        **tt_qty,
                    })

                # Dòng tổng cộng
                tong_row_ch = {
                    "Mã hàng": "", "Tên hàng": "TỔNG CỘNG",
                    "ĐVT": "", "Đơn giá": "",
                    "Tổng SL":    sum(tong_tt_ch.values()),
                    "Thành tiền": f"{int(tong_tien_ncc):,}",
                }
                for tt in tt_co_hang:
                    tong_row_ch[tt] = tong_tt_ch[tt] or ""
                rows_ch.append(tong_row_ch)

                tong_tien_str = f"{int(tong_tien_ncc):,} đ"
                with st.expander(
                    f"🚚 {ncc_name} — {sp_count} sản phẩm | "
                    f"{len(tt_co_hang)} Trung Tâm | Tổng tiền: {tong_tien_str}",
                    expanded=False
                ):
                    st.dataframe(
                        pd.DataFrame(rows_ch),
                        use_container_width=True,
                        hide_index=True,
                    )

            st.divider()
            branch_list_ch = list(all_branches_chuyen.keys())
            selected_ch = st.multiselect(
                "Chọn trung tâm muốn tạo phiếu",
                options=branch_list_ch,
                default=branch_list_ch,
                key="ms_chuyen"
            )

            if selected_ch and st.button("🖨️ Tạo phiếu chuyển hàng từ Excel", type="primary", use_container_width=True):
                from docxcompose.composer import Composer
                now_ch      = datetime.combine(ngay_chuyen_inp, datetime.now().time())
                ngay_str_ch = now_ch.strftime("%d/%m/%Y %H:%M")
                count_ch    = 0

                progress_ch = st.progress(0, text="Đang tạo phiếu...")

                # ── Tạo bảng kê đầu file ──────────────────────────────────
                summary_ch = Document()
                sec_ch = summary_ch.sections[0]
                sec_ch.page_width    = Cm(21)
                sec_ch.page_height   = Cm(29.7)
                sec_ch.top_margin    = Cm(1.5)
                sec_ch.bottom_margin = Cm(1.5)
                sec_ch.left_margin   = Cm(2)
                sec_ch.right_margin  = Cm(1.5)
                sec_ch.header.is_linked_to_previous = False
                sec_ch.footer.is_linked_to_previous = False
                for p in sec_ch.header.paragraphs: p.clear()
                for p in sec_ch.footer.paragraphs: p.clear()

                p_title_ch = summary_ch.add_paragraph()
                p_title_ch.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_title_ch.paragraph_format.space_before = Pt(20)
                p_title_ch.paragraph_format.space_after  = Pt(4)
                r_title_ch = p_title_ch.add_run("BẢNG KÊ TỔNG HỢP PHIẾU CHUYỂN HÀNG")
                r_title_ch.bold = True
                r_title_ch.font.size = Pt(16)
                r_title_ch.font.name = "Arial"
                r_title_ch.font.color.rgb = RGBColor(43, 84, 126)

                p_date_ch = summary_ch.add_paragraph()
                p_date_ch.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_date_ch.paragraph_format.space_after = Pt(16)
                r_date_ch = p_date_ch.add_run(f"Ngày: {now_ch.strftime('%d/%m/%Y')}")
                r_date_ch.font.size = Pt(11)
                r_date_ch.font.name = "Arial"

                col_w_ch = [Cm(1.2), Cm(5.5), Cm(5.5), Cm(4)]
                tbl_ch = summary_ch.add_table(rows=1, cols=4)
                tbl_ch.style = "Table Grid"
                for cell, text, w in zip(tbl_ch.rows[0].cells,
                                         ["STT", "Trung Tâm", "Mã phiếu", "Số lượng SP"],
                                         col_w_ch):
                    cell.width = w
                    set_cell_bg(cell, "FFFFFF")
                    set_cell_border(cell, "000000")
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after  = Pt(3)
                    run = p.add_run(text)
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.name = "Arial"

                # Build tasks trước để điền bảng kê
                tasks_ch = []
                for i, branch in enumerate(selected_ch):
                    raw_items = all_branches_chuyen[branch]
                    items_ch_tmp = [
                        {
                            "ma_hang":   it.get("Mã hàng", ""),
                            "ten_hang":  it.get("Tên Hàng", ""),
                            "dvt":       it.get("Cách đóng", ""),
                            "sl_chuyen": it.get("Số đặt", 0),
                        }
                        for it in raw_items if it.get("Số đặt", 0) > 0
                    ]
                    if not items_ch_tmp:
                        continue
                    ma_phieu_ch = f"PCH-{now_ch.strftime('%d%m%y%H%M%S')}-{i+1:03d}"
                    tasks_ch.append((branch, items_ch_tmp, ma_phieu_ch))

                grand_total_ch = 0
                for i, (branch, items_ch_tmp, ma_phieu_ch) in enumerate(tasks_ch):
                    tong_sp = sum(it["sl_chuyen"] for it in items_ch_tmp)
                    grand_total_ch += tong_sp
                    shade = "F0F4F8" if i % 2 == 0 else "FFFFFF"
                    row_tbl = tbl_ch.add_row()
                    for cell, text, w, align in zip(
                        row_tbl.cells,
                        [str(i+1), branch, ma_phieu_ch, fmt_number(tong_sp)],
                        col_w_ch,
                        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                         WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER]
                    ):
                        cell.width = w
                        set_cell_bg(cell, shade)
                        set_cell_border(cell, "CCCCCC")
                        p = cell.paragraphs[0]
                        p.alignment = align
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.space_after  = Pt(2)
                        run = p.add_run(text)
                        run.font.size = Pt(10)
                        run.font.name = "Arial"

                row_total_ch = tbl_ch.add_row()
                for cell, text, w, align in zip(
                    row_total_ch.cells,
                    ["", "TỔNG CỘNG", "", fmt_number(grand_total_ch)],
                    col_w_ch,
                    [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                     WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER]
                ):
                    cell.width = w
                    set_cell_bg(cell, "EEEEEE")
                    set_cell_border(cell, "000000")
                    p = cell.paragraphs[0]
                    p.alignment = align
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after  = Pt(3)
                    run = p.add_run(text)
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.name = "Arial"

                # Section break sau bảng kê
                p_sep_ch = summary_ch.add_paragraph()
                p_sep_ch.paragraph_format.space_before = Pt(0)
                p_sep_ch.paragraph_format.space_after  = Pt(0)
                pPr_ch = p_sep_ch._p.get_or_add_pPr()
                sectPr_ch = OxmlElement("w:sectPr")
                pgSz_ch = OxmlElement("w:pgSz")
                pgSz_ch.set(qn("w:w"), "12240")
                pgSz_ch.set(qn("w:h"), "16838")
                sectPr_ch.append(pgSz_ch)
                pPr_ch.append(sectPr_ch)

                base_doc_ch = None
                composer_ch = None

                for i, (branch, items_ch_tmp, ma_phieu_ch) in enumerate(tasks_ch):
                    docx_bytes_ch = build_phieu_chuyen(
                        ma_phieu         = ma_phieu_ch,
                        ngay_chuyen      = ngay_str_ch,
                        chi_nhanh_chuyen = chi_nhanh_chuyen_inp,
                        chi_nhanh_nhan   = branch,
                        ghi_chu_chuyen   = ghi_chu_chuyen_inp,
                        ghi_chu_nhan     = ghi_chu_nhan_inp,
                        items            = items_ch_tmp,
                        logo_bytes       = logo_bytes_tab3,
                    )
                    doc_ch = Document(io.BytesIO(docx_bytes_ch))

                    if base_doc_ch is None:
                        base_doc_ch = doc_ch
                        composer_ch = Composer(base_doc_ch)
                    else:
                        composer_ch.append(doc_ch)

                    count_ch += 1
                    progress_ch.progress((i + 1) / len(selected_ch),
                                         text=f"Đang tạo: {branch} ({i+1}/{len(selected_ch)})")

                progress_ch.empty()

                # Ghép bảng kê vào đầu
                final_composer_ch = Composer(summary_ch)
                final_composer_ch.append(base_doc_ch)

                buf_ch = io.BytesIO()
                final_composer_ch.save(buf_ch)
                buf_ch.seek(0)

                st.success(f"✅ Đã tạo xong **{count_ch} phiếu chuyển hàng**!")
                
                # Lưu thông tin vào session để hỏi sau
                st.session_state["pending_sheet_excel"] = {
                    "tasks_ch":         tasks_ch,
                    "ngay_str_ch":      ngay_str_ch,
                    "chi_nhanh_chuyen": chi_nhanh_chuyen_inp,
                    "buf_ch":           buf_ch.getvalue(),   
                    "basename":         chuyen_basename,        
                    "count_ch":         count_ch,
                }
                st.rerun()

        # Hiện hộp xác nhận nếu có phiếu chờ lưu
        if "pending_sheet_excel" in st.session_state:
            p_ex = st.session_state["pending_sheet_excel"]
            st.success(f"✅ Đã tạo xong **{p_ex.get('count_ch', '')} phiếu chuyển hàng**!")
            col_dl_ex, col_yes2, col_no2 = st.columns([3, 1.5, 1])
            
            col_dl_ex.download_button(
                label     = f"⬇️ Tải về {p_ex.get('basename', 'phieu_chuyen')}.docx",
                data      = p_ex["buf_ch"],
                file_name = f"{p_ex.get('basename', 'phieu_chuyen')}.docx",
                mime      = "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key       = "dl_excel_phieu",
            )
            col_yes2.button("📋 Lưu lên Sheet", use_container_width=True, key="yes_sheet_excel")
            col_no2.button("✖ Bỏ qua", use_container_width=True, key="no_sheet_excel")

            if st.session_state.get("yes_sheet_excel"):
                try:
                    sheet = get_gsheet()
                    all_rows = []
                    highlight_indices = []
                    row_counter = 0

                    for branch, items_ch_tmp, ma_phieu_ch in p_ex["tasks_ch"]:
                        items_ch_tmp = merge_items_by_ma_hang(items_ch_tmp)  # ← THÊM DÒNG NÀY
                        highlight_indices.append(row_counter)
                        for item in items_ch_tmp:
                            all_rows.append([
                                ma_phieu_ch, p_ex["ngay_str_ch"], p_ex["chi_nhanh_chuyen"], branch,
                                item.get("ten_hang", ""), item.get("ma_hang", ""),
                                item.get("sl_chuyen", 0), "", item.get("ghi_chu_gop", ""),
                            ])
                            row_counter += 1

                    if all_rows:
                        sheet.insert_rows(all_rows, row=2, value_input_option="USER_ENTERED")

                        sheet_id = sheet.id
                        requests = [
                            # Reset toàn bộ về trắng
                            {
                                "repeatCell": {
                                    "range": {"sheetId": sheet_id, "startRowIndex": 1, "endRowIndex": 1 + len(all_rows), "startColumnIndex": 0, "endColumnIndex": 10},
                                    "cell": {"userEnteredFormat": {"backgroundColor": {"red": 1, "green": 1, "blue": 1}, "textFormat": {"bold": False}}},
                                    "fields": "userEnteredFormat(backgroundColor,textFormat)"
                                }
                            }
                        ]
                        # Thêm highlight cho từng dòng đầu TT
                        for idx in highlight_indices:
                            row_num = 1 + idx  # 0-indexed
                            requests.append({
                                "repeatCell": {
                                    "range": {"sheetId": sheet_id, "startRowIndex": row_num, "endRowIndex": row_num + 1, "startColumnIndex": 0, "endColumnIndex": 10},
                                    "cell": {"userEnteredFormat": {"backgroundColor": {"red": 0.68, "green": 0.85, "blue": 0.90}, "textFormat": {"bold": True}}},
                                    "fields": "userEnteredFormat(backgroundColor,textFormat)"
                                }
                            })

                        # Gửi tất cả format 1 lần duy nhất
                        sheet.spreadsheet.batch_update({"requests": requests})

                    st.success("✅ Đã lưu tất cả phiếu lên Google Sheet!")
                except Exception as e:
                    st.error(f"Lỗi: {e}")
                del st.session_state["pending_sheet_excel"]

            if st.session_state.get("no_sheet_excel"):
                del st.session_state["pending_sheet_excel"]
                st.rerun()

    st.divider()
# CHUYỂN ĐỔI ĐỊNH DẠNG PHIẾU HÀNG KID
elif page == "hang_phong_choi":
    st.subheader("🎮 Hàng Phòng Chơi")
    st.caption("Upload file order → Tự động tách theo NCC → Tải về từng file Excel chuẩn")

    col_up_hpc, col_link_hpc = st.columns([1, 1])
    with col_up_hpc:
        uploaded_hpc = st.file_uploader("Chọn file Excel order", type=["xlsx"], key="hpc_uploader")
    with col_link_hpc:
        link_hpc = st.text_input("Hoặc nhập link Google Sheet (public)", key="link_hpc", placeholder="https://docs.google.com/spreadsheets/d/...")

    file_bytes_hpc = None
    if uploaded_hpc:
        file_bytes_hpc = uploaded_hpc.read()
    elif link_hpc:
        try:
            with st.spinner("Đang tải dữ liệu từ Google Sheet..."):
                file_bytes_hpc = read_gsheet_as_rows(link_hpc)
        except Exception as e:
            st.error(f"Lỗi tải Google Sheet: {e}")

    if file_bytes_hpc:
        import openpyxl
        wb_raw = openpyxl.load_workbook(io.BytesIO(file_bytes_hpc), data_only=True)
        ws_raw = wb_raw.active
        rows_raw = list(ws_raw.iter_rows(values_only=True))

        row1 = rows_raw[1]  # tên trung tâm
        row2 = rows_raw[2]  # sub-header

        # Tìm cột cố định từ hàng 2
        col_ncc      = next((i for i, v in enumerate(row2) if v and str(v).strip() == "NCC"), 2)
        col_ten_hang = next((i for i, v in enumerate(row2) if v and "tên hàng" in str(v).strip().lower()), 3)
        col_dvt      = next((i for i, v in enumerate(row2) if v and str(v).strip().upper() == "ĐVT"), 5)
        col_don_gia  = next((i for i, v in enumerate(row2) if v and "đơn giá" in str(v).strip().lower()), 7)

        # Tìm trung tâm → cột Số đặt
        branches = {}
        for i, val in enumerate(row1):
            if not val or not str(val).strip() or i <= col_don_gia:
                continue
            ten_tt = str(val).strip()
            if ten_tt in ("Tổng order", "None"):
                continue
            # Xác nhận cột này là "Số đặt" ở hàng 2
            if i < len(row2) and row2[i] and "đặt" in str(row2[i]).lower() and "tổng" not in str(row2[i]).lower():
                branches[ten_tt] = i

        # Đọc dữ liệu từ hàng 5 (index 4)
        products = []
        for row in rows_raw[4:]:
            if not row or not row[col_ten_hang]:
                continue
            ten_hang = str(row[col_ten_hang]).strip()
            if not ten_hang or "Tổng" in ten_hang:
                continue
            ncc = str(row[col_ncc]).strip() if row[col_ncc] else ""
            if not ncc or ncc.lower() == "none":
                continue
            dvt = str(row[col_dvt]).strip() if row[col_dvt] else ""
            try:
                don_gia = int(float(row[col_don_gia])) if row[col_don_gia] else 0
            except:
                don_gia = 0

            branch_qty = {}
            for tt, col in branches.items():
                if col < len(row) and row[col]:
                    try:
                        qty = float(row[col])
                        if qty > 0:
                            branch_qty[tt] = qty
                    except:
                        pass

            if branch_qty:
                products.append({
                    "ncc": ncc,
                    "ten_hang": ten_hang,
                    "dvt": dvt,
                    "don_gia": don_gia,
                    "branch_qty": branch_qty,
                })

        if not products:
            st.warning("Không tìm thấy dữ liệu hợp lệ.")
        else:
            # Gom theo NCC
            from collections import defaultdict
            ncc_map = defaultdict(list)
            for p in products:
                ncc_map[p["ncc"]].append(p)

            all_tt = list(branches.keys())

            tong_tat_ca = sum(
                round(sum(p["branch_qty"].values()) * p["don_gia"], 2)
                for items in ncc_map.values() for p in items
            )
            tong_tat_ca_str = f"{tong_tat_ca:,}" + " đ"

            st.success(f"Tìm thấy **{len(ncc_map)} NCC** | **{len(products)} sản phẩm** | **{len(branches)} trung tâm**")
            st.markdown(f"""
            <div style='background:#2B547E; color:white; padding:12px 20px; border-radius:8px;
                        font-size:16px; font-weight:700; text-align:left; margin-top:4px; margin-bottom:16px;'>
                💰 Tổng tiền tất cả NCC: {tong_tat_ca_str}
            </div>
            """, unsafe_allow_html=True)

            # ── Hiển thị report ───────────────────────────────────────────
            import pandas as pd
            all_tt = list(branches.keys())

            for ncc_name, items in ncc_map.items():
                # Chỉ lấy trung tâm có hàng của NCC này
                tt_co_hang = [
                    tt for tt in all_tt
                    if any(p["branch_qty"].get(tt, 0) > 0 for p in items)
                ]

                report_rows = []
                tong_tien_ncc = 0
                tong_tt = {tt: 0 for tt in tt_co_hang}

                for p in items:
                    tong_sp = sum(p["branch_qty"].get(tt, 0) for tt in tt_co_hang)
                    tien_sp = tong_sp * p["don_gia"]
                    tong_tien_ncc += tien_sp

                    row_data = {
                        "Tên hàng":    p["ten_hang"],
                        "ĐVT":         p["dvt"],
                        "Đơn giá":     p["don_gia"],
                    }
                    for tt in tt_co_hang:
                        qty = p["branch_qty"].get(tt, 0)
                        row_data[tt] = qty
                        tong_tt[tt] += qty
                    report_rows.append(row_data)

                # Dòng tổng
                tong_row = {"Tên hàng": "TỔNG", "ĐVT": "", "Đơn giá": ""}
                for tt in tt_co_hang:
                    tong_row[tt] = tong_tt[tt]
                report_rows.append(tong_row)

                df = pd.DataFrame(report_rows)

                tong_tien_str = f"{tong_tien_ncc:,}" + " đ"
                with st.expander(f"🏭 {ncc_name} — {len(items)} sản phẩm | {len(tt_co_hang)} Trung Tâm | Tổng tiền: {tong_tien_str}", expanded=False):
                    st.dataframe(df, use_container_width=True, hide_index=True)

            # ── Xuất file Excel theo từng NCC ─────────────────────────────
            st.divider()
            st.markdown("#### ⬇️ Tải file Excel chuẩn theo NCC")

            ncc_list = list(ncc_map.keys())
            selected_ncc = st.multiselect(
                "Chọn NCC muốn xuất",
                options=ncc_list,
                default=ncc_list,
                key="ms_ncc_hpc"
            )

            if selected_ncc and st.button("📥 Xuất file Excel", type="primary", use_container_width=True):
                all_tt = list(branches.keys())

                from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
                from openpyxl.utils import get_column_letter

                hdr_font   = Font(name="Arial", bold=True, color="FFFFFF", size=10)
                hdr_fill   = PatternFill("solid", fgColor="2B547E")
                tt_fill    = PatternFill("solid", fgColor="1F4E79")
                total_fill = PatternFill("solid", fgColor="1F4E79")
                center     = Alignment(horizontal="center", vertical="center", wrap_text=True)
                left_al    = Alignment(horizontal="left",   vertical="center")
                right_al   = Alignment(horizontal="right",  vertical="center")
                thin       = Side(style="thin", color="CCCCCC")
                bdr        = Border(left=thin, right=thin, top=thin, bottom=thin)

                for ncc_name in selected_ncc:
                    items_ncc = ncc_map[ncc_name]
                    wb_out = openpyxl.Workbook()
                    ws = wb_out.active
                    ws.title = ncc_name[:31]

                    # Cấu trúc cột:
                    # A=NCC, B=Mã hàng, C=Tên hàng, D=ĐVT, E=Đơn giá
                    # F=Số đặt tổng, G=Thành tiền tổng (không có tên TT hàng 1)
                    # H trở đi: mỗi TT 2 cột (Số đặt + Thành tiền), tên TT merge hàng 1

                    TT_COL_START = 8  # cột H
                    tt_co_hang = [
                        tt for tt in all_tt          # ← giữ nguyên all_tt gốc
                        if any(p["branch_qty"].get(tt, 0) > 0 for p in items_ncc)
                    ]

                    ws.cell(1, 6).value = "Tổng order"
                    ws.cell(1, 6).font = Font(name="Arial", bold=True, color="FFFFFF", size=10)
                    ws.cell(1, 6).fill = hdr_fill
                    ws.cell(1, 6).alignment = center
                    ws.cell(1, 6).border = bdr
                    # ── Hàng 1: Tên Trung Tâm từ cột H, merge 2 cột mỗi TT ──
                    for i, tt in enumerate(tt_co_hang):
                        col = TT_COL_START + i * 2
                        cell = ws.cell(1, col)
                        cell.value = tt
                        cell.font = Font(name="Arial", bold=True, color="FFFFFF", size=10)
                        cell.fill = tt_fill
                        cell.alignment = center
                        ws.merge_cells(start_row=1, start_column=col, end_row=1, end_column=col + 1)

                    # ── Hàng 2: Sub-header đầy đủ ──
                    # A-G cố định, rồi lặp Số đặt/Thành tiền cho từng TT
                    row2_vals = ["NCC", "Mã hàng", "Tên hàng", "ĐVT", "Đơn giá", "Số đặt", "Thành Tiền"]
                    for _ in tt_co_hang:
                        row2_vals += ["Số đặt", "Thành tiền"]

                    for c, val in enumerate(row2_vals, 1):
                        cell = ws.cell(2, c)
                        cell.value = val
                        cell.font = hdr_font
                        cell.fill = hdr_fill
                        cell.alignment = center
                        cell.border = bdr

                    # ── Hàng 3: Tổng cộng dùng công thức SUM ──
                    last_row = 3 + len(items_ncc)  # hàng cuối dữ liệu

                    ws.cell(3, 3).font = Font(name="Arial", bold=True, size=10, color="000000")
                    ws.cell(3, 4).value = 0

                    # Cột F3: SUM các cột số đặt của từng TT (H4, J4, L4,... trở xuống)
                    tt_qty_cols = [get_column_letter(TT_COL_START + i * 2) for i in range(len(tt_co_hang))]
                    sum_parts_f = "+".join([f"SUM({col}4:{col}{last_row})" for col in tt_qty_cols])
                    cell_f3 = ws.cell(3, 6)
                    cell_f3.value = f"={sum_parts_f}"
                    cell_f3.font = Font(name="Arial", bold=True, size=10, color="000000")
                    cell_f3.fill = total_fill
                    cell_f3.alignment = center
                    cell_f3.border = bdr

                    # Cột G3: SUM cột thành tiền tổng
                    cell_g3 = ws.cell(3, 7)
                    cell_g3.value = f"=SUM(G4:G{last_row})"
                    cell_g3.font = Font(name="Arial", bold=True, size=10, color="000000")
                    cell_g3.fill = total_fill
                    cell_g3.alignment = center
                    cell_g3.border = bdr

                    # Từng TT: SUM số đặt và SUM thành tiền
                    for i, tt in enumerate(tt_co_hang):
                        col = TT_COL_START + i * 2
                        col_letter_qty = get_column_letter(col)
                        col_letter_amt = get_column_letter(col + 1)

                        cell_qty = ws.cell(3, col)
                        cell_qty.value = f"=SUM({col_letter_qty}4:{col_letter_qty}{last_row})"
                        cell_qty.font = Font(name="Arial", bold=True, size=10, color="000000")
                        cell_qty.fill = total_fill
                        cell_qty.alignment = center
                        cell_qty.border = bdr

                        cell_amt = ws.cell(3, col + 1)
                        cell_amt.value = f"=SUM({col_letter_amt}4:{col_letter_amt}{last_row})"
                        cell_amt.font = Font(name="Arial", bold=True, size=10, color="000000")
                        cell_amt.fill = total_fill
                        cell_amt.alignment = center
                        cell_amt.border = bdr

                    # ── Hàng 4+: Dữ liệu dùng công thức ──
                    for row_idx, p in enumerate(items_ncc, 4):
                        shade = "F5F9FF" if row_idx % 2 == 0 else "FFFFFF"
                        rfill = PatternFill("solid", fgColor=shade)

                        fixed_vals = [
                            (1, p["ncc"],             center),
                            (2, p.get("ma_hang", ""), left_al),
                            (3, p["ten_hang"],        left_al),
                            (4, p["dvt"],             center),
                            (5, p["don_gia"],         right_al),
                        ]
                        for c, val, aln in fixed_vals:
                            cell = ws.cell(row_idx, c)
                            cell.value = val
                            cell.font = Font(name="Arial", size=10, color="000000")
                            cell.fill = rfill
                            cell.alignment = aln
                            cell.border = bdr

                        # Cột F: SUMIFS — cộng tất cả số đặt các TT của hàng này (cột H, J, L,...)
                        tt_qty_cols = [get_column_letter(TT_COL_START + i * 2) for i in range(len(tt_co_hang))]
                        sumifs_parts = "+".join([f"{col}{row_idx}" for col in tt_qty_cols])
                        cell_f = ws.cell(row_idx, 6)
                        cell_f.value = f"={sumifs_parts}"
                        cell_f.font = Font(name="Arial", size=10, color="000000")
                        cell_f.fill = rfill
                        cell_f.alignment = center
                        cell_f.border = bdr

                        # Cột G: Thành tiền tổng = Số đặt tổng × Đơn giá
                        cell_g = ws.cell(row_idx, 7)
                        cell_g.value = f"=F{row_idx}*E{row_idx}"
                        cell_g.font = Font(name="Arial", size=10, color="000000")
                        cell_g.fill = rfill
                        cell_g.alignment = right_al
                        cell_g.border = bdr

                        # Từng TT: số đặt (giá trị) + thành tiền = số đặt × đơn giá
                        for i, tt in enumerate(tt_co_hang):
                            col = TT_COL_START + i * 2
                            qty = p["branch_qty"].get(tt, 0)
                            col_letter = get_column_letter(col)

                            cell_qty = ws.cell(row_idx, col)
                            cell_qty.value = qty or None
                            cell_qty.font = Font(name="Arial", size=10, color="000000")
                            cell_qty.fill = rfill
                            cell_qty.alignment = center
                            cell_qty.border = bdr

                            cell_amt = ws.cell(row_idx, col + 1)
                            # Thành tiền TT = số đặt TT × đơn giá (cột E)
                            cell_amt.value = f"={col_letter}{row_idx}*E{row_idx}" if qty else None
                            cell_amt.font = Font(name="Arial", size=10, color="000000")
                            cell_amt.fill = rfill
                            cell_amt.alignment = center
                            cell_amt.border = bdr

                    # ── Độ rộng cột ──
                    ws.column_dimensions["A"].width = 14
                    ws.column_dimensions["B"].width = 14
                    ws.column_dimensions["C"].width = 35
                    ws.column_dimensions["D"].width = 8
                    ws.column_dimensions["E"].width = 12
                    ws.column_dimensions["F"].width = 10
                    ws.column_dimensions["G"].width = 14
                    for i in range(len(tt_co_hang) * 2):
                        ws.column_dimensions[get_column_letter(TT_COL_START + i)].width = 11

                    ws.row_dimensions[1].height = 30
                    ws.row_dimensions[2].height = 25
                    ws.freeze_panes = "F3"

                    buf_out = io.BytesIO()
                    wb_out.save(buf_out)
                    buf_out.seek(0)

                    safe_ncc = ncc_name.strip().replace(" ", "_").replace("/", "-")
                    st.download_button(
                        label=f"⬇️ {ncc_name}",
                        data=buf_out,
                        file_name=f"{safe_ncc}_chuan.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        key=f"dl_{safe_ncc}",
                        use_container_width=True,
                    )
# CHUYỂN ĐỔI ĐỊNH DẠNG PHIẾU HÀNG GAME
elif page == "hang_game":
    st.subheader("🕹️ Hàng Game")
    st.caption("Upload file Game Excel → Tự động tách theo NCC → Tải về từng file Excel chuẩn")

    col_up_game, col_link_game = st.columns([1, 1])
    with col_up_game:
        uploaded_game = st.file_uploader("Chọn file Excel Game", type=["xlsx"], key="game_uploader")
    with col_link_game:
        link_game = st.text_input("Hoặc nhập link Google Sheet (public)", key="link_game", placeholder="https://docs.google.com/spreadsheets/d/...")

    file_bytes_game = None
    if uploaded_game:
        file_bytes_game = uploaded_game.read()
    elif link_game:
        try:
            with st.spinner("Đang tải dữ liệu từ Google Sheet..."):
                file_bytes_game = read_gsheet_as_rows(link_game)
        except Exception as e:
            st.error(f"Lỗi tải Google Sheet: {e}")

    if file_bytes_game:
        wb_raw = openpyxl.load_workbook(io.BytesIO(file_bytes_game), data_only=True)
        ws_raw = wb_raw.active

        # Hàng 1: tiêu đề chung (bỏ qua)
        # Hàng 2: tên trung tâm
        # Hàng 3: sub-header (Giá cost, Đơn giá/Cái, Số đặt,...)
        # Hàng 4: Tổng Cộng (bỏ qua)
        # Hàng 5+: dữ liệu

        row2 = [ws_raw.cell(2, c).value for c in range(1, ws_raw.max_column + 1)]
        row3 = [ws_raw.cell(3, c).value for c in range(1, ws_raw.max_column + 1)]

        # Cột cố định (1-indexed)
        COL_NCC      = 4   # D
        COL_TEN_HANG = 6   # F
        COL_DVT      = 10  # J (Đơn vị đặt hàng)
        COL_DON_GIA  = 9   # I (Giá cost)
        COL_QUY_CACH = 8   # H (Quy cách đóng gói)

        # Tìm tất cả cột "Số đặt" trong hàng 3
        so_dat_cols = [i + 1 for i, v in enumerate(row3)
                       if v and str(v).strip().lower() == "số đặt"]

        # Tìm trung tâm: hàng 2 từ col 14 trở đi
        # Bỏ qua các giá trị không phải tên TT
        SKIP_KEYWORDS = {"tháng", "tổng order", "none", "1.8", ""}
        branches = {}
        for i, val in enumerate(row2):
            if i + 1 < 14:  # bỏ các cột header cố định
                continue
            if not val or not isinstance(val, str):
                continue
            val_strip = val.strip()
            if not val_strip or val_strip.lower() in SKIP_KEYWORDS:
                continue
            # Chỉ lấy nếu là tên TT thực sự (không chứa keyword header)
            skip = False
            for kw in ["nhóm", "mã hàng", "phụ trách", "ncc", "vat", "tên hàng",
                       "ảnh", "quy cách", "tổng", "tháng"]:
                if kw in val_strip.lower():
                    skip = True
                    break
            if skip:
                continue
            try:
                float(val_strip)
                continue  # bỏ số
            except ValueError:
                pass

            col_1idx = i + 1
            # Số theo quy cách = cột TT + 5
            nearest_so_dat = col_1idx + 5

            # Tránh trùng tên TT
            unique_name = val_strip
            count = sum(1 for k in branches if k == val_strip or k.startswith(val_strip + "_"))
            if count > 0:
                unique_name = f"{val_strip}_{count + 1}"
            branches[unique_name] = nearest_so_dat

        # Đọc dữ liệu từ hàng 5
        products = []
        for row_idx in range(5, ws_raw.max_row + 1):
            ten_hang = ws_raw.cell(row_idx, COL_TEN_HANG).value
            if not ten_hang:
                continue
            ten_hang = str(ten_hang).strip()
            if not ten_hang or "tổng" in ten_hang.lower():
                continue

            ncc = str(ws_raw.cell(row_idx, COL_NCC).value or "").strip()
            if not ncc or ncc.lower() in ("none", ""):
                continue

            dvt = "set"
            try:
                don_gia = float(ws_raw.cell(row_idx, COL_DON_GIA).value or 0)
            except (ValueError, TypeError):
                don_gia = 0

            ma_hang      = str(ws_raw.cell(row_idx, 2).value or "").strip()
            quy_cach     = str(ws_raw.cell(row_idx, COL_QUY_CACH).value or "").strip()
            don_vi_dat   = str(ws_raw.cell(row_idx, COL_DVT).value or "").strip()
            quy_cach_str = str(int(float(quy_cach))) if quy_cach and str(quy_cach).replace('.','',1).isdigit() else quy_cach
            if quy_cach_str and don_vi_dat:
                parts = f"{quy_cach_str} {don_vi_dat}/set"
            elif quy_cach_str:
                parts = f"{quy_cach_str}/set"
            elif don_vi_dat:
                parts = don_vi_dat
            else:
                parts = ""
            ten_hang_full = f"{ten_hang} ({parts})" if parts else ten_hang

            branch_qty = {}
            for tt, so_dat_col in branches.items():
                try:
                    raw = ws_raw.cell(row_idx, so_dat_col).value
                    qty = float(raw or 0)
                    if qty > 0:
                        branch_qty[tt] = qty
                except (ValueError, TypeError):
                    pass

            if branch_qty:
                products.append({
                    "ncc":        ncc,
                    "ma_hang":    ma_hang,
                    "ten_hang":   ten_hang_full,
                    "dvt":        dvt,
                    "don_gia":    don_gia,
                    "branch_qty": branch_qty,
                })

        ncc_map = {}
        for p in products:
            ncc_map.setdefault(p["ncc"], []).append(p)

        if not products:
            st.warning("Không tìm thấy dữ liệu hợp lệ.")
            st.stop()

        all_tt = list(branches.keys())

        tong_tat_ca = sum(
            sum(p["branch_qty"].get(tt, 0) for tt in all_tt) * p["don_gia"]
            for items in ncc_map.values() for p in items
        )
        tong_tat_ca_str = f"{tong_tat_ca:,}" + " đ"

        st.success(f"Tìm thấy **{len(ncc_map)} NCC** | **{len(products)} sản phẩm** | **{len(all_tt)} trung tâm**")
        st.markdown(f"""
        <div style='background:#2B547E; color:white; padding:12px 20px; border-radius:8px;
                    font-size:16px; font-weight:700; margin-top:4px; margin-bottom:16px;'>
            💰 Tổng tiền tất cả NCC: {tong_tat_ca_str}
        </div>
        """, unsafe_allow_html=True)

        import pandas as pd
        for ncc_name, items in ncc_map.items():
            tt_co_hang = [tt for tt in all_tt if any(p["branch_qty"].get(tt, 0) > 0 for p in items)]
            report_rows = []
            tong_tien_ncc = 0
            tong_tt = {tt: 0 for tt in tt_co_hang}
            for p in items:
                tong_sp = sum(p["branch_qty"].get(tt, 0) for tt in tt_co_hang)
                tong_tien_ncc += round(tong_sp * p["don_gia"], 2)
                row_data = {"Tên hàng": p["ten_hang"], "ĐVT": p["dvt"], "Đơn giá": p["don_gia"]}
                for tt in tt_co_hang:
                    qty = p["branch_qty"].get(tt, 0)
                    row_data[tt] = round(qty) if qty else 0
                    tong_tt[tt] += qty
                report_rows.append(row_data)
            tong_row = {"Tên hàng": "TỔNG", "ĐVT": "", "Đơn giá": ""}
            for tt in tt_co_hang:
                tong_row[tt] = tong_tt[tt]
            report_rows.append(tong_row)
            tong_tien_str = f"{tong_tien_ncc:,}" + " đ"
            with st.expander(f"🏭 {ncc_name} — {len(items)} sản phẩm | {len(tt_co_hang)} Trung Tâm | Tổng tiền: {tong_tien_str}", expanded=False):
                st.dataframe(pd.DataFrame(report_rows), use_container_width=True, hide_index=True)

        st.divider()
        st.markdown("#### ⬇️ Tải file Excel chuẩn theo NCC")
        ncc_list = list(ncc_map.keys())
        selected_ncc = st.multiselect("Chọn NCC muốn xuất", options=ncc_list, default=ncc_list, key="ms_ncc_game")

        if selected_ncc and st.button("📥 Xuất file Excel", type="primary", use_container_width=True, key="btn_export_game"):
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from openpyxl.utils import get_column_letter

            hdr_font   = Font(name="Arial", bold=True, color="FFFFFF", size=10)
            hdr_fill   = PatternFill("solid", fgColor="2B547E")
            tt_fill    = PatternFill("solid", fgColor="1F4E79")
            total_fill = PatternFill("solid", fgColor="1F4E79")
            center     = Alignment(horizontal="center", vertical="center", wrap_text=True)
            left_al    = Alignment(horizontal="left",   vertical="center")
            right_al   = Alignment(horizontal="right",  vertical="center")
            thin       = Side(style="thin", color="CCCCCC")
            bdr        = Border(left=thin, right=thin, top=thin, bottom=thin)
            TT_COL_START = 8

            for ncc_name in selected_ncc:
                items_ncc  = ncc_map[ncc_name]
                wb_out     = openpyxl.Workbook()
                ws         = wb_out.active
                ws.title = ncc_name[:31].replace("/", "-").replace("\\", "-").replace("?", "").replace("*", "").replace("[", "").replace("]", "").replace(":", "")
                tt_co_hang = [tt for tt in all_tt if any(p["branch_qty"].get(tt, 0) > 0 for p in items_ncc)]
                last_row   = 3 + len(items_ncc)

                ws.cell(1, 6).value     = "Tổng order"
                ws.cell(1, 6).font      = Font(name="Arial", bold=True, color="FFFFFF", size=10)
                ws.cell(1, 6).fill      = hdr_fill
                ws.cell(1, 6).alignment = center
                ws.cell(1, 6).border    = bdr
                for i, tt in enumerate(tt_co_hang):
                    col  = TT_COL_START + i * 2
                    cell = ws.cell(1, col)
                    cell.value = tt
                    cell.font  = Font(name="Arial", bold=True, color="FFFFFF", size=10)
                    cell.fill  = tt_fill
                    cell.alignment = center
                    cell.border    = bdr
                    ws.merge_cells(start_row=1, start_column=col, end_row=1, end_column=col + 1)

                row2_vals = ["NCC", "Mã hàng", "Tên hàng", "ĐVT", "Đơn giá", "Số đặt", "Thành Tiền"]
                for _ in tt_co_hang:
                    row2_vals += ["Số đặt", "Thành tiền"]
                for c, val in enumerate(row2_vals, 1):
                    cell = ws.cell(2, c)
                    cell.value = val
                    cell.font  = hdr_font
                    cell.fill  = hdr_fill
                    cell.alignment = center
                    cell.border    = bdr

                ws.cell(3, 3).font      = Font(name="Arial", bold=True, size=10)
                ws.cell(3, 3).fill      = total_fill
                ws.cell(3, 3).alignment = center
                ws.cell(3, 3).border    = bdr

                tt_qty_cols = [get_column_letter(TT_COL_START + i * 2) for i in range(len(tt_co_hang))]
                cell_f3 = ws.cell(3, 6)
                cell_f3.value     = "=" + "+".join([f"SUM({c}4:{c}{last_row})" for c in tt_qty_cols])
                cell_f3.font      = Font(name="Arial", bold=True, size=10)
                cell_f3.fill      = total_fill
                cell_f3.alignment = center
                cell_f3.border    = bdr

                cell_g3 = ws.cell(3, 7)
                cell_g3.value     = f"=SUM(G4:G{last_row})"
                cell_g3.font      = Font(name="Arial", bold=True, size=10)
                cell_g3.fill      = total_fill
                cell_g3.alignment = center
                cell_g3.border    = bdr

                for i, tt in enumerate(tt_co_hang):
                    col = TT_COL_START + i * 2
                    cl  = get_column_letter(col)
                    ca  = get_column_letter(col + 1)
                    for c, formula in [(col, f"=SUM({cl}4:{cl}{last_row})"), (col + 1, f"=SUM({ca}4:{ca}{last_row})")]:
                        cell = ws.cell(3, c)
                        cell.value = formula
                        cell.font  = Font(name="Arial", bold=True, size=10)
                        cell.fill  = total_fill
                        cell.alignment = center
                        cell.border    = bdr

                for row_idx, p in enumerate(items_ncc, 4):
                    rfill = PatternFill("solid", fgColor="F5F9FF" if row_idx % 2 == 0 else "FFFFFF")
                    for c, val, aln in [
                        (1, p["ncc"], center), (2, p["ma_hang"], left_al),
                        (3, p["ten_hang"], left_al), (4, p["dvt"], center), (5, p["don_gia"], right_al),
                    ]:
                        cell = ws.cell(row_idx, c)
                        cell.value = val
                        cell.font  = Font(name="Arial", size=10)
                        cell.fill  = rfill
                        cell.alignment = aln
                        cell.border    = bdr

                    tt_cols = [get_column_letter(TT_COL_START + i * 2) for i in range(len(tt_co_hang))]
                    cell_f  = ws.cell(row_idx, 6)
                    cell_f.value     = "=" + "+".join([f"{c}{row_idx}" for c in tt_cols]) if tt_cols else 0
                    cell_f.font      = Font(name="Arial", size=10)
                    cell_f.fill      = rfill
                    cell_f.alignment = center
                    cell_f.border    = bdr

                    cell_g = ws.cell(row_idx, 7)
                    cell_g.value     = f"=F{row_idx}*E{row_idx}"
                    cell_g.font      = Font(name="Arial", size=10)
                    cell_g.fill      = rfill
                    cell_g.alignment = right_al
                    cell_g.border    = bdr

                    for i, tt in enumerate(tt_co_hang):
                        col   = TT_COL_START + i * 2
                        col_l = get_column_letter(col)
                        qty   = p["branch_qty"].get(tt, 0)
                        cq    = ws.cell(row_idx, col)
                        cq.value = round(qty) if qty else None  # hiển thị số chẵn
                        cq.font  = Font(name="Arial", size=10)
                        cq.fill  = rfill
                        cq.alignment = center
                        cq.border    = bdr
                        ca    = ws.cell(row_idx, col + 1)
                        ca.value = round(qty * p["don_gia"], 2) if qty else None
                        ca.font  = Font(name="Arial", size=10)
                        ca.fill  = rfill
                        ca.alignment = center
                        ca.border    = bdr

                for col_letter, width in [("A",14),("B",14),("C",35),("D",8),("E",12),("F",10),("G",14)]:
                    ws.column_dimensions[col_letter].width = width
                for i in range(len(tt_co_hang) * 2):
                    ws.column_dimensions[get_column_letter(TT_COL_START + i)].width = 11
                ws.row_dimensions[1].height = 30
                ws.row_dimensions[2].height = 25
                ws.freeze_panes = "F3"

                buf_out = io.BytesIO()
                wb_out.save(buf_out)
                buf_out.seek(0)
                safe_ncc = ncc_name.strip().replace(" ", "_").replace("/", "-")
                st.download_button(
                    label=f"⬇️ {ncc_name}",
                    data=buf_out,
                    file_name=f"{safe_ncc}_game.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key=f"dl_game_{safe_ncc}",
                    use_container_width=True,
                )
